"""
优化版Word文档格式修改
1. 所有字体设置都使用中英文混合字体
2. 完全通过样式进行批量格式修改
3. 为一级标题添加奇数页分节符
4. 设置标题大纲层级，支持Word目录生成
5. 自动页码设置（页脚居中）
6. 自动页眉设置（奇偶页不同）
7. 自动添加可更新的目录
"""

import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import qn
from docx.oxml import OxmlElement


def set_mixed_font_for_style(style, chinese_font='宋体', english_font='Times New Roman'):
    """为样式设置中英文混合字体"""
    try:
        font = style.font
        
        # 设置基础字体（英文）
        font.name = english_font
        
        # 设置东亚字体（中文）
        rPr = font._element.get_or_add_rPr()
        rfonts = rPr.get_or_add_rFonts()
        rfonts.set(qn('w:ascii'), english_font)      # ASCII字符（英文）
        rfonts.set(qn('w:hAnsi'), english_font)      # 高位ASCII字符
        rfonts.set(qn('w:eastAsia'), chinese_font)   # 东亚字符（中文）
        rfonts.set(qn('w:cs'), english_font)         # 复杂脚本字符
        
        return True
    except Exception as e:
        print(f"    设置样式混合字体失败: {e}")
        return False

def create_optimized_styles(doc):
    """创建优化的格式样式（所有样式都使用混合字体）"""
    
    print("🎨 创建优化格式样式（混合字体）...")
    
    styles_created = {}
    
    try:
        # 1. 文档标题样式
        if 'Title' not in [s.name for s in doc.styles]:
            title_style = doc.styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
            
            # 混合字体设置：中文黑体，英文Times New Roman，二号(22pt)
            set_mixed_font_for_style(title_style, '黑体', 'Times New Roman')
            title_style.font.size = Pt(22)  # 二号
            title_style.font.bold = True
            
            # 段落设置：居中，行距固定20磅，段后24磅
            para_format = title_style.paragraph_format
            para_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para_format.line_spacing = Pt(20)
            para_format.space_before = Pt(0)
            para_format.space_after = Pt(24)
            
            styles_created['title'] = 'Title'
            print("  ✅ 文档标题样式: 中文黑体+英文Times New Roman 22pt，居中")
        
        # 2. 一级标题样式
        if 'Heading1' not in [s.name for s in doc.styles]:
            heading1_style = doc.styles.add_style('Heading1', WD_STYLE_TYPE.PARAGRAPH)
            
            # 混合字体设置：中文黑体，英文Times New Roman，三号(16pt)
            set_mixed_font_for_style(heading1_style, '黑体', 'Times New Roman')
            heading1_style.font.size = Pt(16)  # 三号
            heading1_style.font.bold = True
            
            # 段落设置：居中，行距固定20磅，段前24磅，段后18磅
            para_format = heading1_style.paragraph_format
            para_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para_format.line_spacing = Pt(20)
            para_format.space_before = Pt(24)
            para_format.space_after = Pt(18)
            
            # 设置大纲级别为1级（用于目录生成）

            pPr = para_format.element.get_or_add_pPr()
            outlineLvl = OxmlElement('w:outlineLvl')
            outlineLvl.set(qn('w:val'), '0')  # 设置大纲级别为1
            pPr.append(outlineLvl)
           
            
            styles_created['heading1'] = 'Heading1'
            print("  ✅ 一级标题样式: 中文黑体+英文Times New Roman 16pt，居中，大纲级别1")
        
        # 3. 二级标题样式
        if 'Heading2' not in [s.name for s in doc.styles]:
            heading2_style = doc.styles.add_style('Heading2', WD_STYLE_TYPE.PARAGRAPH)
            
            # 混合字体设置：中文宋体，英文Times New Roman，小三号(15pt)
            set_mixed_font_for_style(heading2_style, '宋体', 'Times New Roman')
            heading2_style.font.size = Pt(15)  # 小三号
            heading2_style.font.bold = True
            
            # 段落设置：不缩进，行距固定20磅，段前18磅，段后12磅
            para_format = heading2_style.paragraph_format
            para_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para_format.line_spacing = Pt(20)
            para_format.space_before = Pt(18)
            para_format.space_after = Pt(12)
            para_format.left_indent = Pt(0)  # 不缩进
            
            # 设置大纲级别为2级（用于目录生成）
            pPr = para_format.element.get_or_add_pPr()
            outlineLvl = OxmlElement('w:outlineLvl')
            outlineLvl.set(qn('w:val'), '1')
            pPr.append(outlineLvl)
            
            styles_created['heading2'] = 'Heading2'
            print("  ✅ 二级标题样式: 中文宋体+英文Times New Roman 15pt加粗，不缩进，大纲级别2")
        
        # 4. 三级标题样式
        if 'Heading3' not in [s.name for s in doc.styles]:
            heading3_style = doc.styles.add_style('Heading3', WD_STYLE_TYPE.PARAGRAPH)
            
            # 混合字体设置：中文宋体，英文Times New Roman，四号(14pt)
            set_mixed_font_for_style(heading3_style, '宋体', 'Times New Roman')
            heading3_style.font.size = Pt(14)  # 四号
            heading3_style.font.bold = True
            
            # 段落设置：缩进2字符，行距固定20磅，段前12磅，段后6磅
            para_format = heading3_style.paragraph_format
            para_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para_format.line_spacing = Pt(20)
            para_format.space_before = Pt(12)
            para_format.space_after = Pt(6)
            para_format.first_line_indent = Pt(28)  # 2字符缩进
            
            # 设置大纲级别为3级（用于目录生成）
            pPr = para_format.element.get_or_add_pPr()
            outlineLvl = OxmlElement('w:outlineLvl')
            outlineLvl.set(qn('w:val'), '2')  # 设置大纲级别为3
            pPr.append(outlineLvl) # 2表示第3级
            
            styles_created['heading3'] = 'Heading3'
            print("  ✅ 三级标题样式: 中文宋体+英文Times New Roman 14pt加粗，缩进2字符，大纲级别3")
        
        # 5. 正文样式
        if 'Body' not in [s.name for s in doc.styles]:
            body_style = doc.styles.add_style('Body', WD_STYLE_TYPE.PARAGRAPH)
            
            # 混合字体设置：中文宋体，英文Times New Roman，小四号(12pt)
            set_mixed_font_for_style(body_style, '宋体', 'Times New Roman')
            body_style.font.size = Pt(12)  # 小四号
            body_style.font.bold = False
            
            # 段落设置：行距固定20磅，段前0磅，段后0磅
            para_format = body_style.paragraph_format
            para_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para_format.line_spacing = Pt(20)
            para_format.space_before = Pt(0)
            para_format.space_after = Pt(0)
            
            styles_created['paragraph'] = 'Body'
            print("  ✅ 正文样式: 中文宋体+英文Times New Roman 12pt，行距20磅")
        
        return styles_created
        
    except Exception as e:
        print(f"❌ 创建优化样式失败: {e}")
        return {}


def add_page_numbers(doc):
    """为所有节添加页码（从第一章开始连续编号）"""
    print(f"\n📄 设置页码...")
    
    def add_page_number_field(footer_run):
        fld_char1 = OxmlElement('w:fldChar')
        fld_char1.set(qn('w:fldCharType'), 'begin')
        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        instr_text.text = "PAGE"
        fld_char2 = OxmlElement('w:fldChar')
        fld_char2.set(qn('w:fldCharType'), 'end')
        footer_run._r.append(fld_char1)
        footer_run._r.append(instr_text)
        footer_run._r.append(fld_char2)
    
    print(len(doc.sections))

    try:
        for section_idx, section in enumerate(doc.sections):
            print(f"  🔧 处理第{section_idx + 1}个分节...")
            footer = section.footer
            footer.is_linked_to_previous = False
            
            # 清空现有页脚内容
            if footer.paragraphs:
                footer_para = footer.paragraphs[0]
                footer_para.clear()
            else:
                footer_para = footer.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 设置页码编号规则
            if section_idx == 0:
                # 目录节：使用希腊字母，从1开始
                sectPr = section._sectPr
                
                # 创建页码设置元素
                pgNumType = sectPr.find(qn('w:pgNumType'))
                if pgNumType is None:
                    pgNumType = OxmlElement('w:pgNumType')
                    sectPr.append(pgNumType)
                
                # 设置罗马数字页码格式，从1开始（i, ii, iii...）
                pgNumType.set(qn('w:start'), '1')
                pgNumType.set(qn('w:fmt'), 'upperRoman')  # 小写罗马数字格式
                
                print(f"      📋 目录节：罗马数字页码从i开始")
            elif section_idx == 1:
                # 第一章：阿拉伯数字重新从1开始
                sectPr = section._sectPr
                
                pgNumType = sectPr.find(qn('w:pgNumType'))
                if pgNumType is None:
                    pgNumType = OxmlElement('w:pgNumType')
                    sectPr.append(pgNumType)
                
                # 设置阿拉伯数字格式，重新从1开始
                pgNumType.set(qn('w:start'), '1')
                pgNumType.set(qn('w:fmt'), 'decimal')
                
                print(f"      🔢 第一章：阿拉伯数字页码重新从1开始")
            else:
                # 后续章节延续阿拉伯数字页码
                print(f"      🔢 第{section_idx}章：页码延续编号")
            
            # 添加页码文本
            prefix_run = footer_para.add_run("-")
            page_run = footer_para.add_run()
            add_page_number_field(page_run)
            suffix_run = footer_para.add_run("-")
            print(f"      ✅ 页码设置完成: '第 [PAGE] 页'")
        
        print(f"\n📊 页码设置总结:")
        print(f"      ✅ 处理分节数: {len(doc.sections)}")
        print(f"      📋 目录节: 罗马数字页码（i, ii, iii...）")
        print(f"      🔢 正文页码: 阿拉伯数字从1开始连续编号")
        return len(doc.sections)
        
    except Exception as e:
        print(f"❌ 设置页码失败: {str(e)}")
        return 0

def add_headers(doc):
    """为所有节添加页眉"""
    
    print(f"\n📄 设置页眉...")
    
    try:
        for section_idx, section in enumerate(doc.sections):
            print(f"  🔧 处理第{section_idx + 1}个分节页眉...")
            
            # 设置奇偶页不同页眉
            section.different_odd_even_pages_header_footer = True
            
            # 目录节页眉
            if section_idx == 0:
                # 奇数页页眉（目录）
                header = section.header
                if header.paragraphs:
                    header_para = header.paragraphs[0]
                    header_para.clear()
                else:
                    header_para = header.add_paragraph()
                
                header_para.text = "目录"
                header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 偶数页页眉（目录）
                even_header = section.even_page_header
                if even_header.paragraphs:
                    even_header_para = even_header.paragraphs[0]
                    even_header_para.clear()
                else:
                    even_header_para = even_header.add_paragraph()
                
                even_header_para.text = "目录"
                even_header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                print(f"      ✅ 目录节页眉设置完成")
                
            else:
                # 正文章节页眉
                # 奇数页页眉（章节标题）
                header = section.header
                if header.paragraphs:
                    header_para = header.paragraphs[0]
                    header_para.clear()
                else:
                    header_para = header.add_paragraph()
                
                header_para.text = f"第{section_idx}章"
                header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 偶数页页眉（文档标题）
                even_header = section.even_page_header
                if even_header.paragraphs:
                    even_header_para = even_header.paragraphs[0]
                    even_header_para.clear()
                else:
                    even_header_para = even_header.add_paragraph()
                
                even_header_para.text = "Word文档格式优化项目"
                even_header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                print(f"      ✅ 第{section_idx}章页眉设置完成")
        
        print(f"\n  📊 页眉设置总结:")
        print(f"      ✅ 处理分节数: {len(doc.sections)}个")
        print(f"      📋 目录节: 奇偶页都显示'目录'")
        print(f"      🔢 正文章节: 奇数页显示章节号，偶数页显示文档标题")
        print(f"      🎯 页眉位置: 页面顶部居中显示")
        
        return len(doc.sections)
        
    except Exception as e:
        print(f"  ❌ 设置页眉失败: {str(e)}")
        import traceback
        print(f"  🔍 错误详情: {traceback.format_exc()}")
        return 0

def add_table_of_contents(doc):
    """在文档开头添加目录并在目录后添加奇数页分节符"""
    
    print(f"\n📋 添加目录...")
    
    try:
        # 创建目录段落
        toc_paragraph = doc.add_paragraph()
        
        # 将目录段落移动到文档开头
        doc._body._body.remove(toc_paragraph._p)
        doc._body._body.insert(0, toc_paragraph._p)
        
        # 添加目录标题
        toc_title = doc.add_paragraph()
        toc_title.text = "目录"
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_title.style.font.size = Pt(16)
        toc_title.style.font.bold = True
        
        # 将目录标题移动到文档开头
        doc._body._body.remove(toc_title._p)
        doc._body._body.insert(0, toc_title._p)
        
        # 创建目录字段
        run = toc_paragraph.add_run()
        
        # 开始字段
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        # 指令文本
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # 1-3级标题，超链接，隐藏页码
        
        # 分隔符
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        # 占位文本
        placeholderText = OxmlElement('w:t')
        placeholderText.text = "请右键点击更新目录"
        
        # 结束字段
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        # 添加到run元素
        r_element = run._r
        r_element.append(fldChar1)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(placeholderText)
        r_element.append(fldChar3)
        
        print(f"  ✅ 目录内容添加成功")
        
        # 在目录段落后添加奇数页分节符
        print(f"  🔧 在目录后添加奇数页分节符...")
        
        # 获取目录段落的XML元素
        toc_para_element = toc_paragraph._p
        
        # 检查是否已有pPr元素
        pPr = toc_para_element.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            toc_para_element.insert(0, pPr)
            print(f"      ➕ 创建pPr元素")
        else:
            print(f"      ✅ 使用现有pPr元素")
        
        # 检查是否已有sectPr元素
        existing_sectPr = pPr.find(qn('w:sectPr'))
        if existing_sectPr is not None:
            pPr.remove(existing_sectPr)
            print(f"      🔄 移除现有sectPr元素")
        
        # 创建新的sectPr元素
        sectPr = OxmlElement('w:sectPr')
        print(f"      ➕ 创建sectPr元素")
        
        # 设置分节符类型为奇数页
        sectType = OxmlElement('w:type')
        sectType.set(qn('w:val'), 'oddPage')
        sectPr.append(sectType)
        print(f"      🎯 设置分节符类型为奇数页")
        
        # 将sectPr添加到pPr中
        pPr.append(sectPr)
        print(f"      ✅ 将sectPr添加到目录段落中")
        
        print(f"  ✅ 目录添加成功")
        print(f"  📋 目录位置: 文档开头")
        print(f"  🔗 包含层级: 1-3级标题")
        print(f"  📄 分节符: 目录后添加奇数页分节符")
        print(f"  💡 使用提示: 在Word中右键点击目录选择'更新域'来刷新目录")
        
        return True
        
    except Exception as e:
        print(f"  ❌ 添加目录失败: {str(e)}")
        import traceback
        print(f"  🔍 错误详情: {traceback.format_exc()}")
        return False

def add_section_breaks(doc, paragraph_types):
    """在每个章节结束位置直接添加奇数页分节符（XML方式）"""
    
    print(f"\n🔧 添加分节符...")
    print(f"📊 文档总段落数: {len(doc.paragraphs)}")
    
    # 第一步：找到所有一级标题的位置
    heading1_positions = []
    for i, para in enumerate(doc.paragraphs):
        para_type = paragraph_types.get(i + 1, 'unknown')
        if para_type == 'heading1':
            heading1_positions.append(i)
            title_text = para.text[:30] + "..." if len(para.text) > 30 else para.text
            print(f"  🔍 发现一级标题: 第{i+1}段 - {title_text}")
    
    print(f"  📋 一级标题总数: {len(heading1_positions)}")
    print(f"  📍 一级标题位置: {[pos+1 for pos in heading1_positions]}")
    
    # 第二步：计算每个章节结束的段落编号
    print(f"\n📖 提取章节结束位置:")
    chapter_end_positions = []
    for i, start_pos in enumerate(heading1_positions):
        if i < len(heading1_positions) - 1:
            # 不是最后一个章节，结束位置是下一个一级标题的前一段
            end_pos = heading1_positions[i + 1] - 1
        else:
            # 最后一个章节，结束位置是文档末尾
            end_pos = len(doc.paragraphs) - 1
        
        chapter_end_positions.append(end_pos)
        chapter_title = doc.paragraphs[start_pos].text[:25] + "..." if len(doc.paragraphs[start_pos].text) > 25 else doc.paragraphs[start_pos].text
        end_text = doc.paragraphs[end_pos].text[:25] + "..." if len(doc.paragraphs[end_pos].text) > 25 else doc.paragraphs[end_pos].text
        
        print(f"  📚 章节{i+1}: {chapter_title}")
        print(f"      起始: 第{start_pos+1}段")
        print(f"      结束: 第{end_pos+1}段 - {end_text}")
    
    print(f"  📋 章节结束位置: {[pos+1 for pos in chapter_end_positions]}")
    
    # 第三步：直接在章节结束段落上设置分节符属性
    print(f"\n🔧 开始设置分节符...")
    print(f"💡 处理策略: 直接在目标段落XML中添加sectPr元素")
    print(f"🎯 需要设置: {len(chapter_end_positions) - 1}个分节符（跳过最后章节）")
    
    sections_added = 0
    
    # 处理除最后一个章节外的所有章节
    for i in range(len(chapter_end_positions) - 1):
        chapter_num = i + 1
        end_pos = chapter_end_positions[i]
        
        print(f"\n  🔧 处理章节{chapter_num}:")
        print(f"      章节结束位置: 第{end_pos+1}段")
        
        try:
            end_para = doc.paragraphs[end_pos]
            end_text = end_para.text[:40] + "..." if len(end_para.text) > 40 else end_para.text
            print(f"      结束段落内容: {end_text}")
            
            # 获取段落的XML元素
            para_element = end_para._p
            print(f"      📄 获取段落XML元素")
            
            # 检查是否已有pPr元素
            pPr = para_element.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                para_element.insert(0, pPr)
                print(f"      ➕ 创建pPr元素")
            else:
                print(f"      ✅ 使用现有pPr元素")
            
            # 检查是否已有sectPr元素
            existing_sectPr = pPr.find(qn('w:sectPr'))
            if existing_sectPr is not None:
                pPr.remove(existing_sectPr)
                print(f"      🔄 移除现有sectPr元素")
            
            # 创建新的sectPr元素
            sectPr = OxmlElement('w:sectPr')
            print(f"      ➕ 创建sectPr元素")
            
            # 设置分节符类型为奇数页
            sectType = OxmlElement('w:type')
            sectType.set(qn('w:val'), 'oddPage')
            sectPr.append(sectType)
            print(f"      🎯 设置分节符类型为奇数页")
            
            # 将sectPr添加到pPr中
            pPr.append(sectPr)
            print(f"      ✅ 将sectPr添加到段落中")
            
            sections_added += 1
            
            print(f"      🎊 章节{chapter_num}分节符设置成功!")
            print(f"      📄 当前文档分节数: {len(doc.sections)}")
            
        except Exception as e:
            print(f"      ❌ 章节{chapter_num}设置分节符失败: {str(e)}")
            import traceback
            print(f"      🔍 错误详情: {traceback.format_exc()}")
    
    print(f"\n📊 分节符设置总结:")
    print(f"  ✅ 成功设置: {sections_added}个分节符")
    print(f"  📄 文档段落数: {len(doc.paragraphs)} (未变化)")
    print(f"  📚 文档分节数: {len(doc.sections)}")
    print(f"  🎯 设置策略: 直接在目标段落XML中添加sectPr元素")
    print(f"  📋 分节符位置: {[chapter_end_positions[i]+1 for i in range(len(chapter_end_positions)-1)]}")
    
    return sections_added

def batch_apply_styles(doc, paragraph_types, styles_created):
    """批量应用样式到文档"""
    
    print(f"\n🔧 批量应用样式...")
    print(f"📊 文档当前段落总数: {len(doc.paragraphs)}")
    print(f"📋 AI分析结果段落数: {len(paragraph_types)}")
    print(f"🎨 可用样式映射: {styles_created}")
    
    # 统计信息
    style_stats = {
        'title': 0,
        'heading1': 0,
        'heading2': 0,
        'heading3': 0,
        'paragraph': 0,
        'unknown': 0
    }
    
    success_count = 0
    
    print(f"\n📖 开始逐段处理...")
    print(f"💡 处理策略: 根据AI分析结果为每段应用对应样式")
    
    # 遍历所有段落并应用样式
    for i, para in enumerate(doc.paragraphs, 1):
        para_type = paragraph_types.get(i, 'unknown')
        text_preview = para.text[:35] + "..." if len(para.text) > 35 else para.text
        
        try:
            if para_type in styles_created:
                # 应用对应的样式
                style_name = styles_created[para_type]
                original_style = para.style.name if para.style else 'None'
                
                para.style = doc.styles[style_name]
                
                style_stats[para_type] += 1
                success_count += 1
                
                # 显示处理进度（每5个段落显示一次或前10个段落）
                if i % 5 == 0 or i <= 10:
                    print(f"  📄 段落{i}: {para_type} → {style_name} | {text_preview}")
                
            else:
                style_stats['unknown'] += 1
                if para_type != 'unknown':
                    print(f"  ⚠️ 段落{i}: 未知类型 {para_type} | {text_preview}")
                
        except Exception as e:
            print(f"  ❌ 段落{i}: 应用样式失败 - {str(e)} | {text_preview}")
            style_stats['unknown'] += 1
    
    print(f"\n📊 样式应用总结:")
    print(f"  ✅ 成功处理: {success_count}个段落")
    print(f"  ❌ 失败/跳过: {len(doc.paragraphs) - success_count}个段落")
    print(f"  📈 成功率: {success_count/len(doc.paragraphs)*100:.1f}%")
    
    print(f"\n📋 各类型段落统计:")
    for style_type, count in style_stats.items():
        if count > 0:
            print(f"  🔸 {style_type}: {count}个")
    
    print(f"\n🎨 样式应用详情:")
    for para_type, style_name in styles_created.items():
        count = style_stats.get(para_type, 0)
        if count > 0:
            print(f"  📝 {para_type} → {style_name}: {count}个段落")
    
    return style_stats, success_count

def load_ai_analysis_result():
    """加载AI分析结果"""
    
    result_file = "backend/ai_analysis_results/ai_analysis_success_20250702_184058.json"
    
    try:
        with open(result_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        analysis_results = data["analysis_result"]["result"]["analysis_result"]
        
        paragraph_types = {}
        for item in analysis_results:
            para_num = item["paragraph_number"]
            para_type = item["type"]
            paragraph_types[para_num] = para_type
        
        print(f"✅ 成功加载AI分析结果: {len(paragraph_types)}个段落")
        return paragraph_types
        
    except Exception as e:
        print(f"❌ 加载AI分析结果失败: {e}")
        return {}

def verify_mixed_fonts(doc):
    """验证混合字体设置"""
    
    print(f"\n🔍 验证混合字体设置...")
    
    try:
        # 检查样式中的字体设置
        custom_styles = [s for s in doc.styles if s.name.startswith('Optimized')]
        
        for style in custom_styles:
            print(f"\n  【{style.name}】")
            
            font = style.font
            print(f"    基础字体: {font.name}")
            print(f"    字体大小: {font.size.pt if font.size else '?'}pt")
            print(f"    粗体: {font.bold}")
            
            # 检查XML中的字体设置
            try:
                rPr = font._element.rPr if font._element.rPr is not None else None
                if rPr is not None:
                    rfonts = rPr.rFonts if hasattr(rPr, 'rFonts') else None
                    if rfonts is not None:
                        ascii_font = rfonts.get(qn('w:ascii'))
                        east_asian_font = rfonts.get(qn('w:eastAsia'))
                        
                        print(f"    ASCII字体(英文): {ascii_font}")
                        print(f"    东亚字体(中文): {east_asian_font}")
                    else:
                        print(f"    字体详情: 无法读取rFonts")
                else:
                    print(f"    字体详情: 无法读取rPr")
            except Exception as e:
                print(f"    字体详情: 读取失败 - {str(e)}")
        
    except Exception as e:
        print(f"❌ 验证混合字体失败: {e}")

def optimized_format_document():
    """优化版文档格式化"""
    
    print("🚀 优化版Word文档格式修改")
    print("=" * 70)
    
    print("📋 优化特性:")
    print("• 所有样式都使用中英文混合字体设置")
    print("• 完全通过样式进行批量格式修改")
    print("• 精确的字体映射：中文字体+英文Times New Roman")
    print("• 统一的格式标准和样式管理")
    print("• 为一级标题添加奇数页分节符")
    print("• 设置标题大纲层级，支持Word目录生成")
    print("• 自动页码设置（页脚居中）")
    print("• 自动页眉设置（奇偶页不同）")
    print("• 自动添加可更新的目录")
    print()
    
    # 1. 加载AI分析结果
    paragraph_types = load_ai_analysis_result()
    if not paragraph_types:
        print("❌ 无法获取段落类型信息，退出")
        return
    
    # 2. 打开原始文档
    source_file = "test_files/test_document.docx"
    
    try:
        doc = Document(source_file)
        print(f"✅ 成功打开文档: {source_file}")
        print(f"📊 文档包含 {len(doc.paragraphs)} 个段落")
    except Exception as e:
        print(f"❌ 打开文档失败: {e}")
        return
    
    # 3. 创建优化样式
    styles_created = create_optimized_styles(doc)
    if not styles_created:
        print("❌ 创建优化样式失败，退出")
        return
    
    # 5. 批量应用样式
    style_stats, success_count = batch_apply_styles(doc, paragraph_types, styles_created)
    
    
    
    # 6.5. 添加分节符
    sections_added = add_section_breaks(doc, paragraph_types)
    
    # 7. 添加目录
    toc_added = add_table_of_contents(doc)

    # 6. 设置页码
    sections_with_page_numbers = add_page_numbers(doc)
    
    # 7. 设置页眉
    sections_with_headers = add_headers(doc)
    
    # 8. 保存优化后的文档
    output_file = "test_files/test_document_optimized.docx"
    
    try:
        doc.save(output_file)
        print(f"\n💾 优化文档已保存: {output_file}")
    except Exception as e:
        print(f"❌ 保存文档失败: {e}")
        return
    

    if toc_added:
        print(f"  目录: 已添加到文档开头")
    
    total_paragraphs = len(doc.paragraphs)
    success_rate = (success_count / total_paragraphs) * 100
    

    if toc_added:
        print(f"📋 目录: 已添加，包含1-3级标题，目录后有奇数页分节符，请在Word中右键更新")

    # 10. 验证混合字体
    verify_mixed_fonts(doc)
    
    print(f"\n🎊 优化完成!")
    print(f"📝 原文件: {source_file}")
    print(f"📝 优化后: {output_file}")


if __name__ == "__main__":
    optimized_format_document()