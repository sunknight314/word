"""
文档格式转换服务 V2
按步骤进行完整的格式转换：
1. 页面设置
2. 创建样式
3. 应用样式
4. 创建章节（分节、页码、页眉页脚）
5. 创建目录
6. 保存文档
"""
import os
from typing import Dict, Any, List, Optional
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from app.services.analysis_result_storage import AnalysisResultStorage
from app.utils.unit_converter import UnitConverter
from app.constants.style_names import StyleNames, get_style_description
import logging

# 设置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentFormatterV2:
    """文档格式化类 V2"""
    
    def __init__(self):
        self.analysis_storage = AnalysisResultStorage()
        self.unit_converter = UnitConverter()
        self.doc = None
        self.format_config = None
        self.analysis_result = None
        
    def format_document(self, source_file_path: str, source_file_id: str, format_file_id: str) -> Dict[str, Any]:
        """
        格式化文档主流程
        """
        try:
            # 准备工作：加载数据
            logger.info("开始文档格式转换...")
            preparation_result = self._prepare_formatting(source_file_path, source_file_id, format_file_id)
            if not preparation_result["success"]:
                return preparation_result
            
            # 打开文档
            self.doc = Document(source_file_path)
            
            # 步骤记录
            steps_results = {}
            
            # 步骤1：应用页面设置
            logger.info("步骤1：应用页面设置")
            page_result = self._step1_apply_page_settings()
            steps_results["page_settings"] = page_result
            
            # 步骤2：创建样式
            logger.info("步骤2：创建样式")
            styles_result = self._step2_create_styles()
            steps_results["create_styles"] = styles_result
            
            # 步骤3：应用样式到段落
            logger.info("步骤3：应用样式")
            apply_result = self._step3_apply_styles()
            steps_results["apply_styles"] = apply_result
            
            # 步骤4：创建章节（分节、页码、页眉页脚）
            logger.info("步骤4：创建章节")
            section_result = self._step4_create_sections()
            steps_results["sections"] = section_result
            
            # 步骤5：创建目录
            logger.info("步骤5：创建目录")
            toc_result = self._step5_create_toc()
            steps_results["toc"] = toc_result
            
            # 步骤6：保存文档
            logger.info("步骤6：保存文档")
            output_path = self._step6_save_document(source_file_path)
            steps_results["save"] = {"success": True, "output_path": output_path}
            
            # 生成处理报告
            report = self._generate_detailed_report(steps_results)
            
            return {
                "success": True,
                "output_path": output_path,
                "steps_results": steps_results,
                "report": report,
                "message": "文档格式化成功完成"
            }
            
        except Exception as e:
            logger.error(f"格式化文档失败: {str(e)}")
            return {
                "success": False,
                "error": f"格式化失败: {str(e)}"
            }
    
    def _prepare_formatting(self, source_file_path: str, source_file_id: str, format_file_id: str) -> Dict[str, Any]:
        """准备格式化所需的数据"""
        # 获取段落分析结果
        paragraph_analysis = self.analysis_storage.get_latest_paragraph_analysis(source_file_id)
        if not paragraph_analysis:
            return {
                "success": False,
                "error": "未找到段落分析结果，请先进行段落分析"
            }
        
        # 获取格式配置
        format_config_data = self.analysis_storage.get_latest_format_config(format_file_id)
        if not format_config_data:
            return {
                "success": False,
                "error": "未找到格式配置，请先生成格式配置"
            }
        
        # 输出数据路径
        logger.info(f"📄 源文档: {source_file_path}")
        if paragraph_analysis.get("file_path"):
            logger.info(f"📊 段落分析: {paragraph_analysis['file_path']}")
        if format_config_data.get("file_path"):
            logger.info(f"🎨 格式配置: {format_config_data['file_path']}")
        
        # 保存到实例变量
        self.format_config = format_config_data.get("format_config", {})
        self.analysis_result = paragraph_analysis.get("result", {}).get("analysis_result", [])
        
        return {"success": True}
    
    def _convert_to_docx_length(self, value: str):
        """将各种单位转换为python-docx需要的长度对象"""
        pt_value = self.unit_converter.convert_to_pt(value)
        if pt_value is not None:
            return Pt(pt_value)
        try:
            return Pt(float(value))
        except:
            return Pt(0)
    
    def _step1_apply_page_settings(self) -> Dict[str, Any]:
        """步骤1：应用页面设置"""
        result = {"success": True, "applied": {}}
        
        page_settings = self.format_config.get("page_settings", {})
        if not page_settings:
            return {"success": True, "applied": {}, "message": "无页面设置配置"}
        
        # 应用边距
        margins = page_settings.get("margins", {})
        if margins:
            for section in self.doc.sections:
                if margins.get("top"):
                    section.top_margin = self._convert_to_docx_length(margins["top"])
                if margins.get("bottom"):
                    section.bottom_margin = self._convert_to_docx_length(margins["bottom"])
                if margins.get("left"):
                    section.left_margin = self._convert_to_docx_length(margins["left"])
                if margins.get("right"):
                    section.right_margin = self._convert_to_docx_length(margins["right"])
            result["applied"]["margins"] = margins
        
        # 应用页面方向和大小
        if page_settings.get("orientation"):
            # TODO: 实现页面方向设置
            result["applied"]["orientation"] = page_settings["orientation"]
        
        if page_settings.get("size"):
            # TODO: 实现页面大小设置
            result["applied"]["size"] = page_settings["size"]
        
        return result
    
    def _step2_create_styles(self) -> Dict[str, Any]:
        """步骤2：创建或更新样式（参考create_styles_from_config方法）"""
        result = {"success": True, "created": [], "updated": [], "failed": [], "styles_mapping": {}}
        
        styles_config = self.format_config.get("styles", {})
        
        for style_key, style_config in styles_config.items():
            try:
                # 为样式名称添加 custom 前缀以区别默认样式
                original_style_name = style_config.get("name", style_key)
                style_name = f"custom{original_style_name}"
                
                # 检查样式是否已存在
                existing_style = None
                try:
                    existing_style = self.doc.styles[style_name]
                    logger.info(f"样式 {style_name} 已存在，将更新配置")
                    result["updated"].append(style_name)
                except KeyError:
                    # 样式不存在，创建新样式
                    existing_style = self.doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                    result["created"].append(style_name)
                    logger.info(f"创建新样式: {style_name}")
                
                # 应用字体设置（增强版，支持中英文混合字体）
                font_config = style_config.get("font", {})
                if font_config:
                    success = self._apply_enhanced_font_settings(existing_style, font_config)
                    if not success:
                        logger.warning(f"样式 {style_name} 字体设置部分失败")
                
                # 应用段落设置（增强版，支持悬挂缩进）
                para_config = style_config.get("paragraph", {})
                if para_config:
                    success = self._apply_enhanced_paragraph_settings(existing_style, para_config)
                    if not success:
                        logger.warning(f"样式 {style_name} 段落设置部分失败")
                
                # 设置大纲级别
                outline_level = style_config.get("outline_level")
                if outline_level is not None:
                    try:
                        self._set_outline_level(existing_style, outline_level)
                        logger.debug(f"样式 {style_name} 设置大纲级别: {outline_level}")
                    except Exception as e:
                        logger.warning(f"设置样式 {style_name} 的大纲级别失败: {str(e)}")
                
                # 记录样式映射
                result["styles_mapping"][style_key] = style_name
                
            except Exception as e:
                result["failed"].append({"style": style_key, "error": str(e)})
                logger.error(f"创建/更新样式 {style_key} 失败: {str(e)}")
        
        return result
    
    def _step3_apply_styles(self) -> Dict[str, Any]:
        """步骤3：应用样式到段落"""
        result = {"success": True, "applied": 0, "skipped": 0, "failed": 0}
        
        # 创建段落号到类型的映射
        para_type_map = {item["paragraph_number"]: item["type"] for item in self.analysis_result}
        styles_config = self.format_config.get("styles", {})
        
        # 遍历文档段落
        for i, paragraph in enumerate(self.doc.paragraphs, 1):
            para_type = para_type_map.get(i)
            if para_type and para_type in styles_config:
                try:
                    # 使用带 custom 前缀的样式名
                    original_style_name = styles_config[para_type].get("name", para_type)
                    style_name = f"custom{original_style_name}"
                    paragraph.style = style_name
                    result["applied"] += 1
                    logger.debug(f"段落 {i} 应用样式: {style_name} (类型: {para_type})")
                except Exception as e:
                    result["failed"] += 1
                    logger.warning(f"段落 {i} 应用样式失败: {str(e)}")
            else:
                result["skipped"] += 1
        
        return result
    
    def _step4_create_sections(self) -> Dict[str, Any]:
        """步骤4：创建章节（处理分节、页码、页眉页脚）"""
        result = {"success": True, "sections_created": 0, "page_numbers": False, "headers_footers": False}
        
        # 处理分节符
        section_breaks = self.format_config.get("section_breaks", {})
        if section_breaks:
            try:
                sections_result = self._add_section_breaks(section_breaks)
                result["sections_created"] = sections_result.get("sections_created", 0)
                logger.info(f"分节符处理完成，创建了 {result['sections_created']} 个分节")
            except Exception as e:
                logger.warning(f"添加分节符失败: {str(e)}")
        
        # 处理页码
        page_numbering = self.format_config.get("page_numbering", {})
        if page_numbering:
            try:
                self._add_page_numbers(page_numbering)
                result["page_numbers"] = True
            except Exception as e:
                logger.warning(f"添加页码失败: {str(e)}")
        
        # 处理页眉页脚
        headers_footers = self.format_config.get("headers_footers", {})
        if headers_footers:
            try:
                self._add_headers_footers(headers_footers)
                result["headers_footers"] = True
            except Exception as e:
                logger.warning(f"添加页眉页脚失败: {str(e)}")
        
        return result
    
    def _step5_create_toc(self) -> Dict[str, Any]:
        """步骤5：创建目录"""
        result = {"success": True, "toc_created": False}
        
        toc_settings = self.format_config.get("toc_settings", {})
        if toc_settings and toc_settings.get("auto_generate_toc", False):
            try:
                # 在文档开头添加目录
                self._insert_toc(toc_settings)
                result["toc_created"] = True
            except Exception as e:
                logger.warning(f"创建目录失败: {str(e)}")
                result["error"] = str(e)
        
        return result
    
    def _step6_save_document(self, source_path: str) -> str:
        """步骤6：保存文档"""
        base_path = Path(source_path)
        output_path = base_path.parent / f"{base_path.stem}_formatted{base_path.suffix}"
        print(f"保存到{output_path}")
        self.doc.save(str(output_path))
        return str(output_path)
    
    def _get_or_create_style(self, style_id: str, style_name: str):
        """获取或创建样式"""
        try:
            return self.doc.styles[style_name]
        except KeyError:
            return self.doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    
    def _apply_font_settings(self, style, font_config: Dict[str, Any]):
        """应用字体设置（旧版方法，保留兼容性）"""
        font = style.font
        
        if font_config.get("size"):
            size_pt = self.unit_converter.convert_to_pt(font_config["size"])
            font.size = Pt(size_pt)
        
        if font_config.get("chinese"):
            font.name = font_config["chinese"]
        
        if "bold" in font_config:
            font.bold = font_config["bold"]
        
        if "italic" in font_config:
            font.italic = font_config["italic"]
    
    def _apply_enhanced_font_settings(self, style, font_config: Dict[str, Any]) -> bool:
        """应用增强版字体设置（支持中英文混合字体）"""
        try:
            font = style.font
            
            # 基本字体设置
            chinese_font = font_config.get("chinese", "宋体")
            english_font = font_config.get("english", "Times New Roman")
            font_size = font_config.get("size", "12pt")
            
            # 设置英文字体为基础字体
            font.name = english_font
            
            # 设置字体大小
            if font_size:
                size_pt = self.unit_converter.convert_to_pt(font_size)
                if size_pt:
                    font.size = Pt(size_pt)
            
            # 设置粗体和斜体
            if "bold" in font_config:
                font.bold = font_config["bold"]
            if "italic" in font_config:
                font.italic = font_config["italic"]
            
            # 设置中英文混合字体
            self._set_mixed_font_for_style(style, chinese_font, english_font)
            
            return True
        except Exception as e:
            logger.error(f"应用增强字体设置失败: {str(e)}")
            return False
    
    def _apply_paragraph_settings(self, style, para_config: Dict[str, Any]):
        """应用段落设置"""
        para_format = style.paragraph_format
        
        # 对齐方式
        alignment = para_config.get("alignment")
        if alignment:
            alignment_map = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
            }
            para_format.alignment = alignment_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
        
        # 行距
        if para_config.get("line_spacing"):
            line_spacing = self.unit_converter.convert_to_pt(para_config["line_spacing"])
            para_format.line_spacing = Pt(line_spacing)
        
        # 段前段后
        if para_config.get("space_before"):
            para_format.space_before = Pt(self.unit_converter.convert_to_pt(para_config["space_before"]))
        if para_config.get("space_after"):
            para_format.space_after = Pt(self.unit_converter.convert_to_pt(para_config["space_after"]))
        
        # 缩进
        if para_config.get("first_line_indent"):
            para_format.first_line_indent = self._convert_to_docx_length(para_config["first_line_indent"])
        if para_config.get("left_indent"):
            para_format.left_indent = self._convert_to_docx_length(para_config["left_indent"])
        if para_config.get("right_indent"):
            para_format.right_indent = self._convert_to_docx_length(para_config["right_indent"])
    
    def _apply_enhanced_paragraph_settings(self, style, para_config: Dict[str, Any]) -> bool:
        """应用增强版段落设置（支持悬挂缩进等高级功能）"""
        try:
            para_format = style.paragraph_format
            
            # 对齐方式
            alignment = para_config.get("alignment", "left")
            alignment_map = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
            }
            para_format.alignment = alignment_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
            
            # 行距
            line_spacing = para_config.get("line_spacing")
            if line_spacing:
                line_spacing_pt = self.unit_converter.convert_to_pt(line_spacing)
                if line_spacing_pt:
                    para_format.line_spacing = Pt(line_spacing_pt)
            
            # 段前距
            space_before = para_config.get("space_before")
            if space_before and space_before != "0pt":
                space_before_pt = self.unit_converter.convert_to_pt(space_before)
                if space_before_pt:
                    para_format.space_before = Pt(space_before_pt)
            
            # 段后距
            space_after = para_config.get("space_after")
            if space_after and space_after != "0pt":
                space_after_pt = self.unit_converter.convert_to_pt(space_after)
                if space_after_pt:
                    para_format.space_after = Pt(space_after_pt)
            
            # 首行缩进
            first_line_indent = para_config.get("first_line_indent")
            if first_line_indent and first_line_indent != "0pt":
                first_line_pt = self.unit_converter.convert_to_pt(first_line_indent)
                if first_line_pt:
                    para_format.first_line_indent = Pt(first_line_pt)
            
            # 左缩进
            left_indent = para_config.get("left_indent")
            if left_indent and left_indent != "0pt":
                left_indent_pt = self.unit_converter.convert_to_pt(left_indent)
                if left_indent_pt:
                    para_format.left_indent = Pt(left_indent_pt)
            
            # 右缩进
            right_indent = para_config.get("right_indent")
            if right_indent and right_indent != "0pt":
                right_indent_pt = self.unit_converter.convert_to_pt(right_indent)
                if right_indent_pt:
                    para_format.right_indent = Pt(right_indent_pt)
            
            # 悬挂缩进（新功能）
            hanging_indent = para_config.get("hanging_indent")
            if hanging_indent and hanging_indent != "0pt":
                hanging_pt = self.unit_converter.convert_to_pt(hanging_indent)
                if hanging_pt:
                    # 悬挂缩进实现为负的首行缩进
                    para_format.first_line_indent = Pt(-hanging_pt)
            
            return True
        except Exception as e:
            logger.error(f"应用增强段落设置失败: {str(e)}")
            return False
    
    def _add_page_numbers(self, page_numbering: Dict[str, Any]):
        """为所有节添加页码"""
        try:
            logger.info("开始设置页码...")
            
            # 获取目录和正文的配置
            toc_config = page_numbering.get("toc_section", {})
            content_config = page_numbering.get("content_sections", {})
            
            # 对齐方式映射
            alignment_map = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT
            }
            
            sections_processed = 0
            for section_idx, section in enumerate(self.doc.sections):
                logger.info(f"处理第{section_idx + 1}个分节的页码...")
                
                # 根据分节类型选择配置
                if section_idx == 0:  # 第一节通常是目录
                    config = toc_config
                    section_type = "目录"
                else:  # 后续节为正文
                    config = content_config  
                    section_type = "正文"
                
                logger.info(f"  分节类型: {section_type}")
                
                # 如果配置为空，跳过该分节
                if not config:
                    logger.info(f"  跳过分节{section_idx + 1}，无页码配置")
                    continue
                
                # 获取该分节的配置
                format_type = config.get("format", "decimal")
                start_number = config.get("start", 1)
                alignment = config.get("alignment", "center")
                template = config.get("template", "{page}")
                location = config.get("location", "footer")
                font_config = config.get("font", {})
                
                logger.info(f"  格式: {format_type}, 起始: {start_number}, 模板: {template}")
                
                # 选择页眉或页脚
                if location == "header":
                    container = section.header
                    logger.info("  位置: 页眉")
                else:
                    container = section.footer
                    logger.info("  位置: 页脚")
                
                container.is_linked_to_previous = False
                
                # 清空现有内容并创建新段落
                if container.paragraphs:
                    container_para = container.paragraphs[0]
                    container_para.clear()
                else:
                    container_para = container.add_paragraph()
                
                # 设置段落对齐
                para_alignment = alignment_map.get(alignment, WD_ALIGN_PARAGRAPH.CENTER)
                container_para.alignment = para_alignment
                
                # 设置页码编号规则
                self._set_page_number_format_by_section(section, section_idx, config)
                
                # 添加页码内容（支持中英文字体）
                self._add_page_number_content_with_font(container_para, template, font_config)
                
                sections_processed += 1
                logger.info(f"第{section_idx + 1}个分节页码设置完成")
            
            logger.info(f"页码设置完成，处理了 {sections_processed} 个分节")
            
        except Exception as e:
            logger.error(f"设置页码失败: {str(e)}")
            raise
    
    def _add_headers_footers(self, headers_footers: Dict[str, Any]):
        """添加页眉页脚"""
        for section in self.doc.sections:
            # 设置奇偶页不同
            if headers_footers.get("odd_even_different"):
                section.different_first_page_header_footer = True
            
            # TODO: 实现具体的页眉页脚内容
            pass
    
    def _insert_toc(self, toc_settings: Dict[str, Any]):
        """插入目录"""
        # 在文档开头插入一个段落作为目录
        toc_para = self.doc.paragraphs[0].insert_paragraph_before()
        toc_para.text = toc_settings.get("title", "目录")
        toc_para.style = "TOCTitle" if "TOCTitle" in self.doc.styles else "Heading1"
        
        # 添加分页符
        from docx.enum.text import WD_BREAK
        toc_para.add_run().add_break(WD_BREAK.PAGE)
        
        # TODO: 实现自动目录生成
        # 这需要使用域代码实现
    
    def _generate_detailed_report(self, steps_results: Dict[str, Any]) -> Dict[str, Any]:
        """生成详细的处理报告"""
        report = {
            "total_steps": 6,
            "completed_steps": sum(1 for step in steps_results.values() if step.get("success", False)),
            "details": steps_results,
            "summary": {
                "page_settings": steps_results.get("page_settings", {}).get("applied", {}),
                "styles_created": len(steps_results.get("create_styles", {}).get("created", [])),
                "styles_updated": len(steps_results.get("create_styles", {}).get("updated", [])),
                "paragraphs_styled": steps_results.get("apply_styles", {}).get("applied", 0),
                "toc_created": steps_results.get("toc", {}).get("toc_created", False),
                "page_numbers_added": steps_results.get("sections", {}).get("page_numbers", False),
                "headers_footers_added": steps_results.get("sections", {}).get("headers_footers", False)
            }
        }
        return report
    
    def _set_outline_level(self, style, outline_level: int):
        """设置样式的大纲级别"""
        try:
            # 获取样式的XML元素
            style_element = style._element
            
            # 查找或创建pPr元素
            pPr = style_element.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                style_element.append(pPr)
            
            # 查找或创建outlineLvl元素
            outline_elem = pPr.find(qn('w:outlineLvl'))
            if outline_elem is None:
                outline_elem = OxmlElement('w:outlineLvl')
                pPr.append(outline_elem)
            
            # 设置大纲级别值
            outline_elem.set(qn('w:val'), str(outline_level))
            
        except Exception as e:
            logger.warning(f"设置大纲级别失败: {str(e)}")
            raise
    
    def _set_mixed_font_for_style(self, style, chinese_font='宋体', english_font='Times New Roman'):
        """为样式设置中英文混合字体"""
        try:
            # 设置基础字体（英文）
            style.font.name = english_font
            
            # 设置东亚字体（中文）
            style_element = style._element
            rpr = style_element.find(qn('w:rPr'))
            if rpr is None:
                rpr = OxmlElement('w:rPr')
                style_element.append(rpr)
            
            # 查找或创建rFonts元素
            rfonts = rpr.find(qn('w:rFonts'))
            if rfonts is None:
                rfonts = OxmlElement('w:rFonts')
                rpr.append(rfonts)
            
            # 设置各种字体类型
            rfonts.set(qn('w:eastAsia'), chinese_font)    # 东亚字体（中文）
            rfonts.set(qn('w:ascii'), english_font)       # ASCII字体（英文）
            rfonts.set(qn('w:hAnsi'), english_font)       # 高ANSI字体（英文）
            
            return True
        except Exception as e:
            logger.error(f"设置混合字体失败: {str(e)}")
            return False
    
    def _add_section_breaks(self, section_breaks_config: Dict[str, Any]) -> Dict[str, Any]:
        """添加分节符，在每个一级标题内容结束时添加分节符"""
        try:
            # 获取一级标题的段落位置
            heading1_paragraphs = []
            for item in self.analysis_result:
                if item.get("type") == "Heading1":
                    heading1_paragraphs.append(item["paragraph_number"] - 1)  # 转为0索引
            
            if not heading1_paragraphs:
                logger.info("未找到一级标题，无需添加分节符")
                return {"success": True, "sections_created": 0}
            
            # 计算每个一级标题内容的结束段落
            section_end_paragraphs = self._calculate_section_end_paragraphs(heading1_paragraphs)
            
            # 获取分节符类型配置
            section_type = section_breaks_config.get("type", "oddPage")
            
            # 添加分节符
            sections_created = self._insert_section_breaks(section_end_paragraphs, section_type)
            
            logger.info(f"成功添加 {sections_created} 个分节符")
            return {"success": True, "sections_created": sections_created}
            
        except Exception as e:
            logger.error(f"添加分节符失败: {str(e)}")
            return {"success": False, "error": str(e)}
    
    def _calculate_section_end_paragraphs(self, heading1_paragraphs: List[int]) -> List[int]:
        """计算每个一级标题内容的结束段落位置"""
        section_end_paragraphs = []
        total_paragraphs = len(self.doc.paragraphs)
        
        for i, start_para in enumerate(heading1_paragraphs):
            if i < len(heading1_paragraphs) - 1:
                # 不是最后一个一级标题：结束位置是下一个一级标题的前一段
                end_para = heading1_paragraphs[i + 1] - 1
                section_end_paragraphs.append(end_para)
            else:
                # 最后一个一级标题：结束位置是文档末尾的前一段
                section_end_paragraphs.append(total_paragraphs - 1)
        
        return section_end_paragraphs
    
    def _insert_section_breaks(self, section_end_paragraphs: List[int], section_type: str) -> int:
        """在指定段落后插入分节符"""
        sections_created = 0
        body = self.doc._body._element
        
        # 从后往前处理，避免段落索引变化的影响
        for end_para_index in reversed(section_end_paragraphs[:-1]):  # 排除最后一节
            try:
                if end_para_index < len(self.doc.paragraphs):
                    # 在指定段落后添加新段落并插入分节符
                    last_para = self.doc.paragraphs[end_para_index]
                    
                    # 创建新段落
                    new_p = OxmlElement('w:p')
                    
                    # 在最后一段后插入新段落
                    last_para._p.addnext(new_p)
                    
                    # 创建分节符
                    sectPr = self._create_sectPr_element(section_type)
                    
                    # 将分节符添加到新段落
                    pPr = new_p.get_or_add_pPr()
                    pPr.append(sectPr)
                    
                    sections_created += 1
                    logger.debug(f"在段落 {end_para_index + 1} 后添加分节符")
                    
            except Exception as e:
                logger.warning(f"在段落 {end_para_index + 1} 后添加分节符失败: {str(e)}")
        
        # 处理最后一节：在文档末尾body层添加分节符
        try:
            existing_sectPr = body.find(qn('w:sectPr'))
            if existing_sectPr is not None:
                body.remove(existing_sectPr)
            
            sectPr = self._create_sectPr_element(section_type, is_last_section=True)
            body.append(sectPr)
            sections_created += 1
            logger.debug("在文档末尾添加分节符")
            
        except Exception as e:
            logger.warning(f"在文档末尾添加分节符失败: {str(e)}")
        
        return sections_created
    
    def _create_sectPr_element(self, section_type: str = "oddPage", is_last_section: bool = False):
        """创建分节符元素"""
        sectPr = OxmlElement('w:sectPr')
        
        # 添加分节符类型
        if section_type:
            type_elem = OxmlElement('w:type')
            type_elem.set(qn('w:val'), section_type)
            sectPr.append(type_elem)
        
        # 添加页面大小 (A4)
        pg_sz = OxmlElement('w:pgSz')
        pg_sz.set(qn('w:w'), '11906')  # 8.27英寸 * 1440
        pg_sz.set(qn('w:h'), '16838')  # 11.69英寸 * 1440
        sectPr.append(pg_sz)
        
        # 添加页边距
        pg_mar = OxmlElement('w:pgMar')
        pg_mar.set(qn('w:top'), '1440')     # 1英寸
        pg_mar.set(qn('w:right'), '1800')   # 1.25英寸
        pg_mar.set(qn('w:bottom'), '1440')  # 1英寸
        pg_mar.set(qn('w:left'), '1800')    # 1.25英寸
        pg_mar.set(qn('w:header'), '851')   # 页眉距离
        pg_mar.set(qn('w:footer'), '992')   # 页脚距离
        pg_mar.set(qn('w:gutter'), '0')     # 装订线
        sectPr.append(pg_mar)
        
        # 添加列设置
        cols = OxmlElement('w:cols')
        cols.set(qn('w:space'), '425')  # 列间距
        cols.set(qn('w:num'), '1')      # 单列
        sectPr.append(cols)
        
        # 添加文档网格
        doc_grid = OxmlElement('w:docGrid')
        doc_grid.set(qn('w:type'), 'lines')
        doc_grid.set(qn('w:linePitch'), '312')
        doc_grid.set(qn('w:charSpace'), '0')
        sectPr.append(doc_grid)
        
        return sectPr
    
    def _set_page_number_format_by_section(self, section, section_idx: int, config: Dict[str, Any]):
        """根据分节配置设置页码格式"""
        try:
            format_type = config.get("format", "decimal")
            start_number = config.get("start", 1)
            restart = config.get("restart", False)
            restart_first_chapter = config.get("restart_first_chapter", False)
            
            sectPr = section._sectPr
            
            # 创建或获取页码设置元素
            pgNumType = sectPr.find(qn('w:pgNumType'))
            if pgNumType is None:
                pgNumType = OxmlElement('w:pgNumType')
                sectPr.append(pgNumType)
            
            # 设置页码格式
            pgNumType.set(qn('w:fmt'), format_type)
            
            # 设置起始页码
            if section_idx == 0:
                # 第一节总是从指定数字开始
                pgNumType.set(qn('w:start'), str(start_number))
                logger.info(f"  第{section_idx + 1}节：{format_type}格式，从{start_number}开始")
            elif section_idx == 1 and restart_first_chapter:
                # 第二节如果设置了重新开始，则重新编号
                pgNumType.set(qn('w:start'), str(start_number))
                logger.info(f"  第{section_idx + 1}节：{format_type}格式，重新从{start_number}开始")
            elif restart:
                # 如果设置了重新开始，则重新编号
                pgNumType.set(qn('w:start'), str(start_number))
                logger.info(f"  第{section_idx + 1}节：{format_type}格式，重新从{start_number}开始")
            else:
                # 继续编号，不设置start属性
                logger.info(f"  第{section_idx + 1}节：{format_type}格式，继续编号")
                
        except Exception as e:
            logger.warning(f"设置第{section_idx + 1}节页码格式失败: {str(e)}")
    
    def _add_page_number_content_with_font(self, container_para, template: str, font_config: Dict[str, Any]):
        """添加页码内容并应用字体设置"""
        try:
            # 解析模板，查找{page}占位符
            parts = template.split('{page}')
            
            if len(parts) == 1:
                # 没有{page}占位符，直接添加文本
                run = container_para.add_run(template)
                self._apply_page_number_font(run, font_config)
            elif len(parts) == 2:
                # 有{page}占位符，分别添加前缀、页码字段、后缀
                if parts[0]:  # 前缀
                    prefix_run = container_para.add_run(parts[0])
                    self._apply_page_number_font(prefix_run, font_config)
                
                # 页码字段
                page_run = container_para.add_run()
                self._add_page_number_field(page_run)
                self._apply_page_number_font(page_run, font_config)
                
                if parts[1]:  # 后缀
                    suffix_run = container_para.add_run(parts[1])
                    self._apply_page_number_font(suffix_run, font_config)
            else:
                # 多个{page}占位符，只处理第一个
                logger.warning("页码模板包含多个{page}占位符，只处理第一个")
                if parts[0]:
                    prefix_run = container_para.add_run(parts[0])
                    self._apply_page_number_font(prefix_run, font_config)
                
                page_run = container_para.add_run()
                self._add_page_number_field(page_run)
                self._apply_page_number_font(page_run, font_config)
                
                remaining = '{page}'.join(parts[1:])
                if remaining:
                    suffix_run = container_para.add_run(remaining)
                    self._apply_page_number_font(suffix_run, font_config)
                    
        except Exception as e:
            logger.warning(f"添加页码内容失败: {str(e)}")
            # 回退到简单的页码显示
            page_run = container_para.add_run()
            self._add_page_number_field(page_run)
            self._apply_page_number_font(page_run, font_config)
    
    def _apply_page_number_font(self, run, font_config: Dict[str, Any]):
        """应用页码字体设置"""
        try:
            if not font_config:
                return
            
            font = run.font
            
            # 设置字体
            chinese_font = font_config.get("chinese")
            english_font = font_config.get("english")
            if english_font:
                font.name = english_font
            
            # 设置字号
            font_size = font_config.get("size")
            if font_size:
                size_pt = self.unit_converter.convert_to_pt(font_size)
                if size_pt:
                    font.size = Pt(size_pt)
            
            # 设置粗体和斜体
            if "bold" in font_config:
                font.bold = font_config["bold"]
            if "italic" in font_config:
                font.italic = font_config["italic"]
            
            # 设置中英文混合字体
            if chinese_font and english_font:
                self._set_mixed_font_for_run(run, chinese_font, english_font)
                
        except Exception as e:
            logger.warning(f"应用页码字体设置失败: {str(e)}")
    
    def _set_mixed_font_for_run(self, run, chinese_font: str, english_font: str):
        """为运行设置中英文混合字体"""
        try:
            # 设置基础字体（英文）
            run.font.name = english_font
            
            # 设置东亚字体（中文）
            run_element = run._r
            rpr = run_element.find(qn('w:rPr'))
            if rpr is None:
                rpr = OxmlElement('w:rPr')
                run_element.append(rpr)
            
            # 查找或创建rFonts元素
            rfonts = rpr.find(qn('w:rFonts'))
            if rfonts is None:
                rfonts = OxmlElement('w:rFonts')
                rpr.append(rfonts)
            
            # 设置各种字体类型
            rfonts.set(qn('w:eastAsia'), chinese_font)    # 东亚字体（中文）
            rfonts.set(qn('w:ascii'), english_font)       # ASCII字体（英文）
            rfonts.set(qn('w:hAnsi'), english_font)       # 高ANSI字体（英文）
            
        except Exception as e:
            logger.warning(f"设置运行混合字体失败: {str(e)}")
    
    def _set_page_number_format(self, section, section_idx: int, format_type: str, start_number: int):
        """设置页码格式和起始编号"""
        try:
            sectPr = section._sectPr
            
            # 创建或获取页码设置元素
            pgNumType = sectPr.find(qn('w:pgNumType'))
            if pgNumType is None:
                pgNumType = OxmlElement('w:pgNumType')
                sectPr.append(pgNumType)
            
            # 根据配置或默认规则设置页码格式
            if section_idx == 0:
                # 第一节（通常是目录）：使用罗马数字
                pgNumType.set(qn('w:start'), str(start_number))
                pgNumType.set(qn('w:fmt'), 'upperRoman')
                logger.info(f"第{section_idx + 1}节：罗马数字页码从{start_number}开始")
            elif section_idx == 1:
                # 第二节（通常是正文第一章）：阿拉伯数字重新从1开始
                pgNumType.set(qn('w:start'), '1')
                pgNumType.set(qn('w:fmt'), 'decimal')
                logger.info(f"第{section_idx + 1}节：阿拉伯数字页码重新从1开始")
            else:
                # 后续章节：延续阿拉伯数字页码
                pgNumType.set(qn('w:fmt'), format_type)
                logger.info(f"第{section_idx + 1}节：页码延续编号")
                
        except Exception as e:
            logger.warning(f"设置第{section_idx + 1}节页码格式失败: {str(e)}")
    
    def _add_page_number_content(self, footer_para, template: str):
        """添加页码内容到页脚段落"""
        try:
            # 解析模板，查找{PAGE}占位符
            parts = template.split('{PAGE}')
            
            if len(parts) == 1:
                # 没有{PAGE}占位符，直接添加文本
                footer_para.add_run(template)
            elif len(parts) == 2:
                # 有{PAGE}占位符，分别添加前缀、页码字段、后缀
                if parts[0]:  # 前缀
                    footer_para.add_run(parts[0])
                
                # 页码字段
                page_run = footer_para.add_run()
                self._add_page_number_field(page_run)
                
                if parts[1]:  # 后缀
                    footer_para.add_run(parts[1])
            else:
                # 多个{PAGE}占位符，只处理第一个
                logger.warning("页码模板包含多个{PAGE}占位符，只处理第一个")
                if parts[0]:
                    footer_para.add_run(parts[0])
                page_run = footer_para.add_run()
                self._add_page_number_field(page_run)
                remaining = '{PAGE}'.join(parts[1:])
                if remaining:
                    footer_para.add_run(remaining)
                    
        except Exception as e:
            logger.warning(f"添加页码内容失败: {str(e)}")
            # 回退到简单的页码显示
            page_run = footer_para.add_run()
            self._add_page_number_field(page_run)
    
    def _add_page_number_field(self, footer_run):
        """添加页码字段到运行中"""
        try:
            # 创建页码字段元素
            fld_char1 = OxmlElement('w:fldChar')
            fld_char1.set(qn('w:fldCharType'), 'begin')
            
            instr_text = OxmlElement('w:instrText')
            instr_text.set(qn('xml:space'), 'preserve')
            instr_text.text = "PAGE"
            
            fld_char2 = OxmlElement('w:fldChar')
            fld_char2.set(qn('w:fldCharType'), 'end')
            
            # 添加到运行元素
            footer_run._r.append(fld_char1)
            footer_run._r.append(instr_text)
            footer_run._r.append(fld_char2)
            
        except Exception as e:
            logger.error(f"添加页码字段失败: {str(e)}")
            raise