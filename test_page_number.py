from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm

# 创建新文档
doc = Document()

# 添加标题和内容
doc.add_heading("测试文档：分节与页码", level=1)
for i in range(1, 6):
    doc.add_paragraph(f"这是第{i}段正文内容。")

# 添加奇数页分节符
doc.add_section(WD_SECTION_START.ODD_PAGE)  # 新节从下一个奇数页开始
new_section = doc.sections[-1]

# 设置新节内容
doc.add_paragraph("\n--- 新节开始（奇数页）---\n")
for i in range(6, 11):
    doc.add_paragraph(f"新节第{i}段内容。")

# ===== 添加页码（所有节） =====
def add_page_number(footer_run):
    """插入页码字段（XML底层操作）"""
    # 创建页码字段元素
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = "PAGE"  # 当前页码
    
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    
    # 组合元素
    footer_run._r.append(fld_char1)
    footer_run._r.append(instr_text)
    footer_run._r.append(fld_char2)

# 为所有节添加页脚页码
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False  # 断开与前节链接
    footer_para = footer.paragraphs[0]
    footer_para.alignment = 1  # 1=居中, 0=左对齐, 2=右对齐
    footer_run = footer_para.add_run()
    add_page_number(footer_run)  # 插入页码

# 保存文档
doc.save("分节页码测试.docx")
print("文档已生成：包含奇数页分节符和统一页码")