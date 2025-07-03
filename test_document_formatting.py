"""
测试python-docx的格式修改功能
根据AI分析结果对test_document.docx进行格式调整
"""

import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import qn

def load_ai_analysis_result():
    """加载AI分析结果"""
    
    result_file = "backend/ai_analysis_results/ai_analysis_success_20250702_184058.json"
    
    try:
        with open(result_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # 提取分析结果
        analysis_results = data["analysis_result"]["result"]["analysis_result"]
        
        # 创建段落类型映射
        paragraph_types = {}
        for item in analysis_results:
            para_num = item["paragraph_number"]
            para_type = item["type"]
            paragraph_types[para_num] = para_type
        
        print(f"✅ 成功加载AI分析结果: {len(paragraph_types)}个段落")
        return paragraph_types
        
    except Exception as e:
        print(f"❌ 加载AI分析结果失败: {e}")
        return {}

def create_format_styles(doc):
    """创建格式样式"""
    
    print("🎨 创建格式样式...")
    
    styles_created = {}
    
    try:
        # 1. 一级标题样式
        if 'Heading1Custom' not in [s.name for s in doc.styles]:
            heading1_style = doc.styles.add_style('Heading1Custom', WD_STYLE_TYPE.PARAGRAPH)
            
            # 字体设置：黑体，三号(16pt)
            font = heading1_style.font
            font.name = '黑体'
            font.size = Pt(16)  # 三号 = 16pt
            font.bold = True
            
            # 段落设置：居中，行距固定20磅，段前24磅，段后18磅
            para_format = heading1_style.paragraph_format
            para_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para_format.line_spacing = Pt(20)  # 固定值20磅
            para_format.space_before = Pt(24)
            para_format.space_after = Pt(18)
            
            styles_created['heading1'] = 'Heading1Custom'
            print("  ✅ 创建一级标题样式: 黑体16pt，居中，行距20磅")
        
        # 2. 二级标题样式
        if 'Heading2Custom' not in [s.name for s in doc.styles]:
            heading2_style = doc.styles.add_style('Heading2Custom', WD_STYLE_TYPE.PARAGRAPH)
            
            # 字体设置：宋体加粗，小三号(15pt)
            font = heading2_style.font
            font.name = '宋体'
            font.size = Pt(15)  # 小三号 = 15pt
            font.bold = True
            
            # 段落设置：不缩进，行距固定20磅，段前18磅，段后12磅
            para_format = heading2_style.paragraph_format
            para_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para_format.line_spacing = Pt(20)
            para_format.space_before = Pt(18)
            para_format.space_after = Pt(12)
            para_format.left_indent = Pt(0)  # 不缩进
            
            styles_created['heading2'] = 'Heading2Custom'
            print("  ✅ 创建二级标题样式: 宋体15pt加粗，不缩进，行距20磅")
        
        # 3. 三级标题样式
        if 'Heading3Custom' not in [s.name for s in doc.styles]:
            heading3_style = doc.styles.add_style('Heading3Custom', WD_STYLE_TYPE.PARAGRAPH)
            
            # 字体设置：宋体，四号(14pt)加粗
            font = heading3_style.font
            font.name = '宋体'
            font.size = Pt(14)  # 四号 = 14pt
            font.bold = True
            
            # 段落设置：缩进2字符，行距固定20磅，段前12磅，段后6磅
            para_format = heading3_style.paragraph_format
            para_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para_format.line_spacing = Pt(20)
            para_format.space_before = Pt(12)
            para_format.space_after = Pt(6)
            para_format.first_line_indent = Pt(28)  # 2字符缩进 ≈ 28磅
            
            styles_created['heading3'] = 'Heading3Custom'
            print("  ✅ 创建三级标题样式: 宋体14pt加粗，缩进2字符，行距20磅")
        
        # 4. 正文样式
        if 'BodyCustom' not in [s.name for s in doc.styles]:
            body_style = doc.styles.add_style('BodyCustom', WD_STYLE_TYPE.PARAGRAPH)
            
            # 字体设置：宋体，小四号(12pt)
            font = body_style.font
            font.name = '宋体'
            font.size = Pt(12)  # 小四号 = 12pt
            
            # 段落设置：行距固定20磅，段前0磅，段后0磅
            para_format = body_style.paragraph_format
            para_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para_format.line_spacing = Pt(20)
            para_format.space_before = Pt(0)
            para_format.space_after = Pt(0)
            
            styles_created['paragraph'] = 'BodyCustom'
            print("  ✅ 创建正文样式: 宋体12pt，行距20磅，无段落间距")
        
        # 5. 文档标题样式
        if 'TitleCustom' not in [s.name for s in doc.styles]:
            title_style = doc.styles.add_style('TitleCustom', WD_STYLE_TYPE.PARAGRAPH)
            
            # 字体设置：黑体，二号(22pt)
            font = title_style.font
            font.name = '黑体'
            font.size = Pt(22)  # 二号 = 22pt
            font.bold = True
            
            # 段落设置：居中，行距固定20磅，段后24磅
            para_format = title_style.paragraph_format
            para_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para_format.line_spacing = Pt(20)
            para_format.space_before = Pt(0)
            para_format.space_after = Pt(24)
            
            styles_created['title'] = 'TitleCustom'
            print("  ✅ 创建文档标题样式: 黑体22pt，居中，行距20磅")
        
        return styles_created
        
    except Exception as e:
        print(f"❌ 创建样式失败: {e}")
        return {}

def set_mixed_font_for_run(run, chinese_font='宋体', english_font='Times New Roman'):
    """为运行设置中英文混合字体"""
    try:
        # 设置基础字体
        run.font.name = english_font
        
        # 设置东亚字体（中文）
        rPr = run._element.get_or_add_rPr()
        rfonts = rPr.get_or_add_rFonts()
        rfonts.set(qn('w:eastAsia'), chinese_font)
        
        return True
    except Exception as e:
        print(f"    设置混合字体失败: {e}")
        return False

def apply_paragraph_formatting(para, para_type, styles_created, doc):
    """应用段落格式"""
    
    try:
        # 根据段落类型应用样式
        if para_type in styles_created:
            style_name = styles_created[para_type]
            para.style = doc.styles[style_name]
            
            # 如果是正文，还需要设置中英文混合字体
            if para_type == 'paragraph':
                for run in para.runs:
                    set_mixed_font_for_run(run, '宋体', 'Times New Roman')
            
            return True
        else:
            print(f"    未知段落类型: {para_type}")
            return False
            
    except Exception as e:
        print(f"    应用格式失败: {e}")
        return False

def format_document():
    """格式化文档"""
    
    print("📄 开始格式化文档...")
    print("=" * 60)
    
    # 1. 加载AI分析结果
    paragraph_types = load_ai_analysis_result()
    
    if not paragraph_types:
        print("❌ 无法获取段落类型信息，退出")
        return
    
    # 2. 打开原始文档
    source_file = "test_files/test_document.docx"
    
    try:
        doc = Document(source_file)
        print(f"✅ 成功打开文档: {source_file}")
        print(f"📊 文档包含 {len(doc.paragraphs)} 个段落")
    except Exception as e:
        print(f"❌ 打开文档失败: {e}")
        return
    
    # 3. 创建自定义样式
    styles_created = create_format_styles(doc)
    
    if not styles_created:
        print("❌ 创建样式失败，退出")
        return
    
    # 4. 应用格式到每个段落
    print(f"\n🔧 应用格式到段落...")
    
    format_stats = {
        'title': 0,
        'heading1': 0,
        'heading2': 0,
        'heading3': 0,
        'paragraph': 0,
        'unknown': 0
    }
    
    for i, para in enumerate(doc.paragraphs, 1):
        # 获取段落类型
        para_type = paragraph_types.get(i, 'unknown')
        
        # 显示段落信息
        text_preview = para.text[:30] + "..." if len(para.text) > 30 else para.text
        print(f"  段落{i}: {para_type} - {text_preview}")
        
        # 应用格式
        if para_type != 'unknown':
            success = apply_paragraph_formatting(para, para_type, styles_created, doc)
            if success:
                format_stats[para_type] += 1
                print(f"    ✅ 格式应用成功")
            else:
                format_stats['unknown'] += 1
                print(f"    ❌ 格式应用失败")
        else:
            format_stats['unknown'] += 1
            print(f"    ⚠️ 跳过未知类型")
    
    # 5. 保存格式化后的文档
    output_file = "test_files/test_document_formatted.docx"
    
    try:
        doc.save(output_file)
        print(f"\n💾 格式化文档已保存: {output_file}")
    except Exception as e:
        print(f"❌ 保存文档失败: {e}")
        return
    
    # 6. 显示格式化统计
    print(f"\n📊 格式化统计:")
    print(f"  文档标题: {format_stats['title']}个")
    print(f"  一级标题: {format_stats['heading1']}个")
    print(f"  二级标题: {format_stats['heading2']}个")
    print(f"  三级标题: {format_stats['heading3']}个")
    print(f"  正文段落: {format_stats['paragraph']}个")
    print(f"  未处理: {format_stats['unknown']}个")
    
    total_processed = sum(format_stats.values()) - format_stats['unknown']
    total_paragraphs = len(doc.paragraphs)
    
    print(f"\n✅ 格式化完成: {total_processed}/{total_paragraphs} ({total_processed/total_paragraphs*100:.1f}%)")

def verify_formatting():
    """验证格式化结果"""
    
    print(f"\n🔍 验证格式化结果...")
    print("=" * 40)
    
    try:
        # 打开格式化后的文档
        formatted_file = "test_files/test_document_formatted.docx"
        doc = Document(formatted_file)
        
        print(f"📄 检查文档: {formatted_file}")
        
        # 检查样式应用情况
        style_usage = {}
        
        for para in doc.paragraphs:
            style_name = para.style.name
            style_usage[style_name] = style_usage.get(style_name, 0) + 1
        
        print(f"\n📈 样式使用统计:")
        for style_name, count in style_usage.items():
            print(f"  {style_name}: {count}个段落")
        
        # 检查几个重要段落的格式
        print(f"\n🔎 重点段落格式检查:")
        
        important_paras = [1, 2, 5, 11]  # 标题、一级标题、二级标题等
        
        for para_num in important_paras:
            if para_num <= len(doc.paragraphs):
                para = doc.paragraphs[para_num - 1]
                style = para.style
                
                print(f"  段落{para_num}: {para.text[:20]}...")
                print(f"    样式: {style.name}")
                print(f"    字体: {style.font.name}")
                print(f"    大小: {style.font.size.pt if style.font.size else '?'}pt")
                print(f"    对齐: {style.paragraph_format.alignment}")
        
        print(f"\n✅ 验证完成，请在Word中打开查看最终效果")
        
    except Exception as e:
        print(f"❌ 验证失败: {e}")

def main():
    """主函数"""
    
    print("🎯 Word文档格式修改测试")
    print("=" * 70)
    
    print("📋 格式要求:")
    print("• 一级标题: 黑体三号(16pt)，居中，行距20磅，段前24磅段后18磅")
    print("• 二级标题: 宋体小三号(15pt)加粗，不缩进，行距20磅，段前18磅段后12磅")
    print("• 三级标题: 宋体四号(14pt)加粗，缩进2字符，行距20磅，段前12磅段后6磅")
    print("• 正文内容: 宋体小四号(12pt)，行距20磅，无段落间距")
    print("• 特殊要求: 正文中文用宋体，英文用Times New Roman")
    print()
    
    # 执行格式化
    format_document()
    
    # 验证结果
    verify_formatting()
    
    print(f"\n🎉 测试完成!")
    print(f"📝 原文件: test_files/test_document.docx")
    print(f"📝 格式化后: test_files/test_document_formatted.docx")
    print(f"💡 在Word中对比两个文件查看格式变化")

if __name__ == "__main__":
    main()