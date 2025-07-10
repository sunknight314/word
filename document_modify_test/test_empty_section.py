#!/usr/bin/env python3
"""
创建一个包含空分节符段落的测试文档
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_test_document():
    """创建包含空分节符段落的测试文档"""
    
    # 创建新文档
    doc = Document()
    
    # 添加一些正常段落
    doc.add_paragraph("第一段：正常内容")
    doc.add_paragraph("第二段：这段后面会有分节符")
    
    # 在第二段添加分节符
    para2 = doc.paragraphs[1]
    p_element = para2._element
    pPr = p_element.get_or_add_pPr()
    
    # 创建分节符
    sectPr = OxmlElement('w:sectPr')
    type_elem = OxmlElement('w:type')
    type_elem.set(qn('w:val'), 'oddPage')
    sectPr.append(type_elem)
    
    # 添加页面设置
    pg_sz = OxmlElement('w:pgSz')
    pg_sz.set(qn('w:w'), '11906')
    pg_sz.set(qn('w:h'), '16838')
    sectPr.append(pg_sz)
    
    pPr.append(sectPr)
    
    # 添加一个空段落并在其中放入分节符
    empty_para = doc.add_paragraph("")  # 空段落
    empty_p_element = empty_para._element
    empty_pPr = empty_p_element.get_or_add_pPr()
    
    # 创建另一个分节符
    sectPr2 = OxmlElement('w:sectPr')
    type_elem2 = OxmlElement('w:type')
    type_elem2.set(qn('w:val'), 'evenPage')
    sectPr2.append(type_elem2)
    
    # 添加页面设置
    pg_sz2 = OxmlElement('w:pgSz')
    pg_sz2.set(qn('w:w'), '11906')
    pg_sz2.set(qn('w:h'), '16838')
    sectPr2.append(pg_sz2)
    
    empty_pPr.append(sectPr2)
    
    # 再添加一些正常段落
    doc.add_paragraph("第四段：正常内容")
    doc.add_paragraph("第五段：最后一段")
    
    # 保存文档
    doc.save("test_with_empty_sections.docx")
    print("✅ 创建测试文档成功: test_with_empty_sections.docx")
    print(f"   - 总段落数: {len(doc.paragraphs)}")
    print(f"   - 第2段包含奇数页分节符（非空段落）")
    print(f"   - 第3段包含偶数页分节符（空段落）")

if __name__ == "__main__":
    create_test_document()