#!/usr/bin/env python3
"""
Word文档样式管理器 - 添加自定义样式到样式库
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import os

class WordStyleManager:
    """Word样式管理器"""
    
    def __init__(self, doc_path):
        """初始化文档"""
        self.doc = Document(doc_path) if os.path.exists(doc_path) else Document()
        self.styles = self.doc.styles
        
    def add_custom_styles(self):
        """添加一套自定义样式"""
        
        print("🎨 开始添加自定义样式...")
        
        # 添加各种类型的样式
        self.add_heading_styles()
        self.add_paragraph_styles() 
        self.add_character_styles()
        self.add_table_styles()
        
        print("✅ 自定义样式添加完成!")
        
    def add_heading_styles(self):
        """添加标题样式"""
        
        print("📝 添加标题样式...")
        
        # 自定义一级标题
        try:
            heading1_style = self.styles.add_style('自定义标题1', WD_STYLE_TYPE.PARAGRAPH)
            heading1_style.base_style = self.styles['Heading 1']
            
            # 字体设置
            font = heading1_style.font
            font.name = '微软雅黑'
            font.size = Pt(18)
            font.bold = True
            font.color.rgb = RGBColor(0, 32, 96)  # 深蓝色
            
            # 段落设置
            paragraph_format = heading1_style.paragraph_format
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph_format.space_before = Pt(24)
            paragraph_format.space_after = Pt(12)
            
            print("  ✅ 添加自定义标题1样式")
            
        except ValueError:
            print("  ⚠️ 自定义标题1样式已存在，跳过")
        
        # 自定义二级标题
        try:
            heading2_style = self.styles.add_style('自定义标题2', WD_STYLE_TYPE.PARAGRAPH)
            heading2_style.base_style = self.styles['Heading 2']
            
            # 字体设置
            font = heading2_style.font
            font.name = '微软雅黑'
            font.size = Pt(16)
            font.bold = True
            font.color.rgb = RGBColor(0, 112, 192)  # 蓝色
            
            # 段落设置
            paragraph_format = heading2_style.paragraph_format
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph_format.space_before = Pt(18)
            paragraph_format.space_after = Pt(6)
            paragraph_format.left_indent = Inches(0)
            
            print("  ✅ 添加自定义标题2样式")
            
        except ValueError:
            print("  ⚠️ 自定义标题2样式已存在，跳过")
            
        # 自定义三级标题
        try:
            heading3_style = self.styles.add_style('自定义标题3', WD_STYLE_TYPE.PARAGRAPH)
            heading3_style.base_style = self.styles['Heading 3']
            
            # 字体设置
            font = heading3_style.font
            font.name = '微软雅黑'
            font.size = Pt(14)
            font.bold = True
            font.color.rgb = RGBColor(68, 114, 196)  # 中蓝色
            
            # 段落设置
            paragraph_format = heading3_style.paragraph_format
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph_format.space_before = Pt(12)
            paragraph_format.space_after = Pt(6)
            
            print("  ✅ 添加自定义标题3样式")
            
        except ValueError:
            print("  ⚠️ 自定义标题3样式已存在，跳过")
    
    def add_paragraph_styles(self):
        """添加段落样式"""
        
        print("📄 添加段落样式...")
        
        # 正文样式
        try:
            body_style = self.styles.add_style('自定义正文', WD_STYLE_TYPE.PARAGRAPH)
            body_style.base_style = self.styles['Normal']
            
            # 字体设置
            font = body_style.font
            font.name = '宋体'
            font.size = Pt(12)
            font.color.rgb = RGBColor(0, 0, 0)  # 黑色
            
            # 段落设置
            paragraph_format = body_style.paragraph_format
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # 两端对齐
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            paragraph_format.line_spacing = 1.5  # 1.5倍行距
            paragraph_format.space_after = Pt(6)
            paragraph_format.first_line_indent = Inches(0.5)  # 首行缩进
            
            print("  ✅ 添加自定义正文样式")
            
        except ValueError:
            print("  ⚠️ 自定义正文样式已存在，跳过")
        
        # 引用样式
        try:
            quote_style = self.styles.add_style('引用文本', WD_STYLE_TYPE.PARAGRAPH)
            quote_style.base_style = self.styles['Normal']
            
            # 字体设置
            font = quote_style.font
            font.name = '楷体'
            font.size = Pt(11)
            font.italic = True
            font.color.rgb = RGBColor(89, 89, 89)  # 深灰色
            
            # 段落设置
            paragraph_format = quote_style.paragraph_format
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph_format.left_indent = Inches(0.5)
            paragraph_format.right_indent = Inches(0.5)
            paragraph_format.space_before = Pt(6)
            paragraph_format.space_after = Pt(6)
            
            print("  ✅ 添加引用文本样式")
            
        except ValueError:
            print("  ⚠️ 引用文本样式已存在，跳过")
        
        # 代码样式
        try:
            code_style = self.styles.add_style('代码块', WD_STYLE_TYPE.PARAGRAPH)
            code_style.base_style = self.styles['Normal']
            
            # 字体设置
            font = code_style.font
            font.name = 'Consolas'
            font.size = Pt(10)
            font.color.rgb = RGBColor(0, 0, 0)
            
            # 段落设置
            paragraph_format = code_style.paragraph_format
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph_format.left_indent = Inches(0.5)
            paragraph_format.space_before = Pt(6)
            paragraph_format.space_after = Pt(6)
            
            print("  ✅ 添加代码块样式")
            
        except ValueError:
            print("  ⚠️ 代码块样式已存在，跳过")
    
    def add_character_styles(self):
        """添加字符样式"""
        
        print("🔤 添加字符样式...")
        
        # 强调文本
        try:
            emphasis_style = self.styles.add_style('强调文本', WD_STYLE_TYPE.CHARACTER)
            
            # 字体设置
            font = emphasis_style.font
            font.bold = True
            font.color.rgb = RGBColor(192, 0, 0)  # 红色
            
            print("  ✅ 添加强调文本样式")
            
        except ValueError:
            print("  ⚠️ 强调文本样式已存在，跳过")
        
        # 关键词
        try:
            keyword_style = self.styles.add_style('关键词', WD_STYLE_TYPE.CHARACTER)
            
            # 字体设置
            font = keyword_style.font
            font.bold = True
            font.color.rgb = RGBColor(0, 176, 80)  # 绿色
            
            print("  ✅ 添加关键词样式")
            
        except ValueError:
            print("  ⚠️ 关键词样式已存在，跳过")
        
        # 代码文本
        try:
            inline_code_style = self.styles.add_style('行内代码', WD_STYLE_TYPE.CHARACTER)
            
            # 字体设置
            font = inline_code_style.font
            font.name = 'Consolas'
            font.size = Pt(11)
            font.color.rgb = RGBColor(199, 37, 78)  # 暗红色
            
            print("  ✅ 添加行内代码样式")
            
        except ValueError:
            print("  ⚠️ 行内代码样式已存在，跳过")
    
    def add_table_styles(self):
        """添加表格样式"""
        
        print("📊 添加表格样式...")
        
        # 注意：表格样式比较复杂，这里创建一个基础的表格样式
        try:
            table_style = self.styles.add_style('自定义表格', WD_STYLE_TYPE.TABLE)
            
            print("  ✅ 添加自定义表格样式")
            
        except ValueError:
            print("  ⚠️ 自定义表格样式已存在，跳过")
    
    def list_all_styles(self):
        """列出文档中的所有样式"""
        
        print("📋 文档中的所有样式:")
        print("=" * 50)
        
        # 按类型分组显示
        style_types = {
            WD_STYLE_TYPE.PARAGRAPH: "段落样式",
            WD_STYLE_TYPE.CHARACTER: "字符样式", 
            WD_STYLE_TYPE.TABLE: "表格样式",
            WD_STYLE_TYPE.LIST: "列表样式"
        }
        
        for style_type, type_name in style_types.items():
            print(f"\n📝 {type_name}:")
            styles_of_type = [s for s in self.styles if s.type == style_type]
            
            for style in styles_of_type:
                builtin_mark = " (内置)" if style.builtin else " (自定义)"
                print(f"  • {style.name}{builtin_mark}")
    
    def apply_styles_demo(self):
        """演示应用样式"""
        
        print("🎨 添加演示内容并应用样式...")
        
        # 添加标题
        title = self.doc.add_paragraph("Word文档样式演示", style='自定义标题1')
        
        # 添加二级标题
        subtitle = self.doc.add_paragraph("1. 段落样式演示", style='自定义标题2')
        
        # 添加正文
        body_text = self.doc.add_paragraph("这是使用自定义正文样式的段落。" +
                                         "这种样式采用宋体字，12磅大小，1.5倍行距，" +
                                         "首行缩进0.5英寸，两端对齐。", style='自定义正文')
        
        # 添加引用
        quote_text = self.doc.add_paragraph("这是一段引用文本，使用楷体，" +
                                          "左右缩进，颜色较浅，用于突出引用内容。", 
                                          style='引用文本')
        
        # 添加三级标题
        subtitle3 = self.doc.add_paragraph("1.1 字符样式演示", style='自定义标题3')
        
        # 添加包含字符样式的段落
        mixed_para = self.doc.add_paragraph("这个段落包含不同的字符样式：", style='自定义正文')
        mixed_para.add_run("强调文本", style='强调文本')
        mixed_para.add_run("、")
        mixed_para.add_run("关键词", style='关键词') 
        mixed_para.add_run("、")
        mixed_para.add_run("print('行内代码')", style='行内代码')
        mixed_para.add_run("等。")
        
        # 添加代码块
        code_para = self.doc.add_paragraph(
            "def hello_world():\n    print('Hello, World!')\n    return True", 
            style='代码块'
        )
        
        print("  ✅ 演示内容添加完成")
    
    def save(self, output_path):
        """保存文档"""
        self.doc.save(output_path)
        print(f"💾 文档已保存: {output_path}")

def main():
    """主函数"""
    
    print("🎯 Word文档样式管理器")
    print("=" * 50)
    
    # 文件路径
    input_file = "../document_modify_test/test_document.docx"
    output_file = "test_document_with_styles.docx"
    
    # 初始化样式管理器
    style_manager = WordStyleManager(input_file)
    
    print(f"📂 处理文档: {input_file}")
    
    # 显示原有样式
    print("\n📋 原有样式:")
    style_manager.list_all_styles()
    
    print("\n" + "="*50)
    
    # 添加自定义样式
    style_manager.add_custom_styles()
    
    print("\n" + "="*50)
    
    # 显示添加后的样式
    print("\n📋 添加样式后:")
    style_manager.list_all_styles()
    
    # 添加演示内容
    print("\n" + "="*50)
    style_manager.apply_styles_demo()
    
    # 保存文档
    style_manager.save(output_file)
    
    print(f"\n🎉 样式添加完成! 请打开 {output_file} 查看效果")

if __name__ == "__main__":
    main()