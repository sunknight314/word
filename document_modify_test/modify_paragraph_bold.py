#!/usr/bin/env python3
"""
使用python-docx修改第8段为粗体
"""

import os
from docx import Document

def make_paragraph_bold(file_path, paragraph_number):
    """将指定段落设置为粗体"""
    
    print(f"🔄 开始修改文档: {file_path}")
    print(f"🎯 目标: 将第 {paragraph_number} 段设置为粗体")
    
    try:
        # 打开文档
        doc = Document(file_path)
        print(f"📄 成功打开文档，共有 {len(doc.paragraphs)} 段")
        
        # 检查段落编号是否有效
        if paragraph_number < 1 or paragraph_number > len(doc.paragraphs):
            print(f"❌ 段落编号 {paragraph_number} 无效，文档只有 {len(doc.paragraphs)} 段")
            return False
        
        # 获取目标段落（段落编号从1开始，数组索引从0开始）
        target_paragraph = doc.paragraphs[paragraph_number - 1]
        
        print(f"📝 第 {paragraph_number} 段内容: {target_paragraph.text[:50]}...")
        
        # 将段落中的所有run设置为粗体
        for run in target_paragraph.runs:
            run.bold = True
            print(f"  ✅ 设置run为粗体: {run.text[:30]}...")
        
        # 如果段落没有runs（空段落或特殊情况），创建一个run
        if not target_paragraph.runs and target_paragraph.text:
            # 为段落添加一个run并设置为粗体
            run = target_paragraph.runs[0] if target_paragraph.runs else target_paragraph.add_run(target_paragraph.text)
            run.bold = True
            print(f"  ✅ 创建并设置新run为粗体")
        
        # 保存文档
        doc.save(file_path)
        print(f"💾 文档保存成功: {file_path}")
        
        return True
        
    except Exception as e:
        print(f"❌ 修改失败: {e}")
        return False

def show_all_paragraphs(file_path):
    """显示文档中所有段落的内容（用于查看）"""
    
    print(f"📋 查看文档段落内容: {file_path}")
    
    try:
        doc = Document(file_path)
        
        print(f"📄 文档共有 {len(doc.paragraphs)} 段:\n")
        
        for i, para in enumerate(doc.paragraphs, 1):
            text = para.text.strip()
            if text:  # 只显示非空段落
                print(f"段落 {i:2d}: {text[:80]}{'...' if len(text) > 80 else ''}")
                
                # 显示该段落的格式信息
                bold_runs = []
                for run in para.runs:
                    if run.bold:
                        bold_runs.append(run.text[:20])
                
                if bold_runs:
                    print(f"         🔥 粗体部分: {', '.join(bold_runs)}")
            else:
                print(f"段落 {i:2d}: (空段落)")
        
        return True
        
    except Exception as e:
        print(f"❌ 查看失败: {e}")
        return False

def main():
    """主函数"""
    
    print("🎯 段落粗体修改器")
    print("=" * 50)
    
    # 文件路径
    file_path = "test_document.docx"
    paragraph_number = 8
    
    # 检查文件是否存在
    if not os.path.exists(file_path):
        print(f"❌ 文件不存在: {file_path}")
        return
    
    # 首先显示所有段落内容
    print("📋 修改前的文档内容:")
    show_all_paragraphs(file_path)
    
    print("\n" + "="*50)
    
    # 修改第8段为粗体
    success = make_paragraph_bold(file_path, paragraph_number)
    
    if success:
        print(f"\n🎉 第 {paragraph_number} 段已成功设置为粗体!")
        
        print("\n📋 修改后的文档内容:")
        show_all_paragraphs(file_path)
    else:
        print(f"\n❌ 修改第 {paragraph_number} 段失败!")

if __name__ == "__main__":
    main()