#!/usr/bin/env python3
"""
根据分节信息自动添加分节符
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def calculate_section_end_paragraphs(start_paragraphs, total_paragraphs):
    """
    根据每节开始段落编号，计算每节的结束段落编号
    
    Args:
        start_paragraphs: 每节开始段落编号列表，如 [0, 3, 7]
        total_paragraphs: 文档总段落数
        
    Returns:
        list: 每节结束段落编号列表，如 [2, 6, "文档末尾"]
    """
    
    end_paragraphs = []
    
    for i, start_para in enumerate(start_paragraphs):
        if i < len(start_paragraphs) - 1:
            # 不是最后一节：结束段落 = 下一节开始段落 - 1
            end_para = start_paragraphs[i + 1] - 1
            end_paragraphs.append(end_para)
        else:
            # 最后一节：到文档末尾
            end_paragraphs.append("文档末尾")
    
    return end_paragraphs

def create_sectPr_element(section_type=None, is_first_section=False):
    """
    创建分节符元素
    
    Args:
        section_type: 分节符类型 ('oddPage', 'evenPage', 'nextPage', 'continuous')
        is_first_section: 是否为第一节（第一节不设置type属性）
        
    Returns:
        XML元素: 分节符元素
    """
    
    sectPr = OxmlElement('w:sectPr')
    
    # 第一节不设置type属性，其他节设置type属性
    if not is_first_section and section_type:
        type_elem = OxmlElement('w:type')
        type_elem.set(qn('w:val'), section_type)
        sectPr.append(type_elem)
    
    # 添加页面大小 (A4)
    pg_sz = OxmlElement('w:pgSz')
    pg_sz.set(qn('w:w'), '11906')  # 8.27英寸 * 1440
    pg_sz.set(qn('w:h'), '16838')  # 11.69英寸 * 1440
    sectPr.append(pg_sz)
    
    # 添加页边距
    pg_mar = OxmlElement('w:pgMar')
    pg_mar.set(qn('w:top'), '1440')     # 1英寸
    pg_mar.set(qn('w:right'), '1800')   # 1.25英寸
    pg_mar.set(qn('w:bottom'), '1440')  # 1英寸
    pg_mar.set(qn('w:left'), '1800')    # 1.25英寸
    pg_mar.set(qn('w:header'), '851')   # 页眉距离
    pg_mar.set(qn('w:footer'), '992')   # 页脚距离
    pg_mar.set(qn('w:gutter'), '0')     # 装订线
    sectPr.append(pg_mar)
    
    # 添加列设置
    cols = OxmlElement('w:cols')
    cols.set(qn('w:space'), '425')  # 列间距
    cols.set(qn('w:num'), '1')      # 单列
    sectPr.append(cols)
    
    # 添加文档网格
    doc_grid = OxmlElement('w:docGrid')
    doc_grid.set(qn('w:type'), 'lines')
    doc_grid.set(qn('w:linePitch'), '312')
    doc_grid.set(qn('w:charSpace'), '0')
    sectPr.append(doc_grid)
    
    return sectPr

def add_sections_by_start_paragraphs(file_path, start_paragraphs, section_types=None, output_path=None):
    """
    根据每节开始段落编号添加分节符
    
    Args:
        file_path: 输入文档路径
        start_paragraphs: 每节开始段落编号列表（从0开始），如 [0, 3, 7]
        section_types: 每节的分节符类型列表，如 ['oddPage', 'evenPage']，第一节会忽略
        output_path: 输出文档路径，如果为None则覆盖原文件
        
    Returns:
        dict: 操作结果
    """
    
    print(f"🔄 开始处理文档: {file_path}")
    print(f"🎯 分节信息: {start_paragraphs}")
    
    try:
        # 打开文档
        doc = Document(file_path)
        body = doc._body._element
        total_paragraphs = len(doc.paragraphs)
        
        print(f"📄 文档总段落数: {total_paragraphs}")
        
        # 计算每节的结束段落
        end_paragraphs = calculate_section_end_paragraphs(start_paragraphs, total_paragraphs)
        
        print("📊 分节计算结果:")
        for i, (start, end) in enumerate(zip(start_paragraphs, end_paragraphs)):
            section_name = f"第{i+1}节"
            if end == "文档末尾":
                print(f"  {section_name}: 段落{start+1} 到 {end}")
            else:
                print(f"  {section_name}: 段落{start+1} 到 段落{end+1}")
        
        print()
        
        # 设置默认分节符类型
        if section_types is None:
            section_types = ['oddPage'] * (len(start_paragraphs) - 1)  # 第一节不需要type
        
        # 从后往前处理，避免段落索引变化
        sections_to_process = list(enumerate(zip(start_paragraphs, end_paragraphs)))
        sections_to_process.reverse()
        
        for section_index, (start_para, end_para) in sections_to_process:
            section_num = section_index + 1
            is_first_section = section_index == 0
            is_last_section = end_para == "文档末尾"
            
            # 获取分节符类型
            if is_first_section:
                section_type = None  # 第一节不设置type
            elif section_index - 1 < len(section_types):
                section_type = section_types[section_index - 1]
            else:
                section_type = 'oddPage'  # 默认奇数页
            
            print(f"🔧 处理第{section_num}节...")
            
            if is_last_section:
                # 最后一节：在文档末尾body层添加分节符
                print(f"  📝 最后一节，在body层添加分节符")
                
                # 检查是否已有body层分节符
                existing_sectPr = body.find(qn('w:sectPr'))
                if existing_sectPr is not None:
                    # 移除现有的
                    body.remove(existing_sectPr)
                
                # 创建新的分节符
                sectPr = create_sectPr_element(section_type, is_first_section)
                body.append(sectPr)
                
                print(f"  ✅ 第{section_num}节处理完成（body层分节符）")
                
            else:
                # 非最后一节：在本节最后一段后添加新段落，分节符放在新段落中
                last_para_index = end_para
                last_para = doc.paragraphs[last_para_index]
                
                print(f"  📝 在第{last_para_index+1}段后添加分节符段落")
                
                # 创建新段落
                new_p = OxmlElement('w:p')
                
                # 在最后一段后插入新段落
                last_para._p.addnext(new_p)
                
                # 创建分节符
                sectPr = create_sectPr_element(section_type, is_first_section)
                
                # 将分节符添加到新段落
                pPr = new_p.get_or_add_pPr()
                pPr.append(sectPr)
                
                type_desc = "无类型" if is_first_section else section_type
                print(f"  ✅ 第{section_num}节处理完成（分节符类型: {type_desc}）")
        
        # 保存文档
        save_path = output_path if output_path else file_path
        doc.save(save_path)
        
        result = {
            'success': True,
            'sections_created': len(start_paragraphs),
            'section_info': list(zip(start_paragraphs, end_paragraphs)),
            'output_path': save_path
        }
        
        print(f"✅ 分节符添加完成!")
        print(f"   - 创建分节数: {len(start_paragraphs)}个")
        print(f"💾 文档已保存: {save_path}")
        
        return result
        
    except Exception as e:
        print(f"❌ 处理失败: {e}")
        return {
            'success': False,
            'error': str(e)
        }

def show_sections_info(file_path):
    """显示文档分节信息"""
    
    print(f"📊 文档分节信息: {file_path}")
    
    try:
        doc = Document(file_path)
        body = doc._body._element
        
        print(f"📄 文档总段落数: {len(doc.paragraphs)}")
        print(f"📖 文档总分节数: {len(doc.sections)}")
        print()
        
        section_count = 0
        
        # 检查段落内的分节符
        for i, para in enumerate(doc.paragraphs):
            p_element = para._element
            pPr = p_element.pPr
            
            if pPr is not None:
                sectPr = pPr.find(qn('w:sectPr'))
                if sectPr is not None:
                    section_count += 1
                    sectType = sectPr.find(qn('w:type'))
                    break_type = sectType.get(qn('w:val')) if sectType is not None else "无类型"
                    
                    para_text = para.text.strip()[:40]
                    if not para_text:
                        para_text = "(空段落-分节符)"
                    else:
                        para_text += "..."
                    
                    print(f"第{section_count}节: 段落{i+1} - {para_text}")
                    print(f"       分节符类型: {break_type}")
                    print()
        
        # 检查body层的分节符
        body_sectPr = body.find(qn('w:sectPr'))
        if body_sectPr is not None:
            section_count += 1
            body_sectType = body_sectPr.find(qn('w:type'))
            break_type = body_sectType.get(qn('w:val')) if body_sectType is not None else "无类型"
            
            print(f"第{section_count}节: 文档末尾 (body层)")
            print(f"       分节符类型: {break_type}")
            print()
        
        print(f"📈 总计: {section_count}个分节")
        
    except Exception as e:
        print(f"❌ 查看失败: {e}")

def main():
    """主函数 - 演示用法"""
    
    print("🎯 Word文档分节符构建器")
    print("=" * 60)
    
    # 示例：将文档分为3节
    # 第1节：段落1-3 (0-2)
    # 第2节：段落4-7 (3-6) 
    # 第3节：段落8-末尾 (7-末尾)
    start_paragraphs = [0, 3, 7]
    section_types = ['oddPage', 'oddPage']  # 第一节不需要指定type
    
    file_path = "cleaned_document.docx"
    output_path = "sectioned_document.docx"
    
    print("📋 处理前的文档信息:")
    show_sections_info(file_path)
    
    print("\n" + "="*60)
    
    # 添加分节符
    result = add_sections_by_start_paragraphs(
        file_path=file_path,
        start_paragraphs=start_paragraphs,
        section_types=section_types,
        output_path=output_path
    )
    
    if result['success']:
        print("\n📋 处理后的文档信息:")
        show_sections_info(output_path)
    else:
        print(f"\n❌ 处理失败: {result['error']}")

if __name__ == "__main__":
    main()