#!/usr/bin/env python3
"""
通用Word分节符处理器
根据Word XML分节符规则：
1. 第一节的w:sectPr没有w:type属性
2. 其他节都可以设置w:type属性
3. 除最后一节外，所有节的w:sectPr都在本节最后一段后新加一段的段内
4. 最后一节的w:sectPr在整个文档末尾body的层级下
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import copy

class WordSectionManager:
    """Word分节符管理器"""
    
    def __init__(self, doc_path):
        """初始化文档"""
        self.doc = Document(doc_path)
        self.body = self.doc._body._element
        
    def get_current_sections_info(self):
        """获取当前分节信息"""
        sections_info = []
        
        # 检查每个段落的分节符
        for i, para in enumerate(self.doc.paragraphs):
            p_element = para._element
            pPr = p_element.pPr
            if pPr is not None:
                sectPr = pPr.find(qn('w:sectPr'))
                if sectPr is not None:
                    # 获取分节符类型
                    sectType = sectPr.find(qn('w:type'))
                    break_type = sectType.get(qn('w:val')) if sectType is not None else None
                    
                    sections_info.append({
                        'paragraph_index': i,
                        'type': break_type,
                        'is_first_section': break_type is None,
                        'sectPr': sectPr
                    })
        
        # 检查文档末尾的分节符（最后一节）
        body_sectPr = self.body.find(qn('w:sectPr'))
        if body_sectPr is not None:
            sectType = body_sectPr.find(qn('w:type'))
            break_type = sectType.get(qn('w:val')) if sectType is not None else None
            
            sections_info.append({
                'paragraph_index': len(self.doc.paragraphs),
                'type': break_type,
                'is_last_section': True,
                'sectPr': body_sectPr
            })
        
        return sections_info
    
    def create_sectPr_element(self, section_type=None, copy_from_sectPr=None):
        """创建分节符元素"""
        
        if copy_from_sectPr is not None:
            # 复制现有的分节符设置
            sectPr = copy.deepcopy(copy_from_sectPr)
            
            # 如果需要设置新的分节符类型
            if section_type is not None:
                # 移除现有的type元素
                existing_type = sectPr.find(qn('w:type'))
                if existing_type is not None:
                    sectPr.remove(existing_type)
                
                # 添加新的type元素
                if section_type != 'none':  # 'none'表示不设置type（第一节）
                    type_elem = OxmlElement('w:type')
                    type_elem.set(qn('w:val'), section_type)
                    sectPr.insert(0, type_elem)
        else:
            # 创建新的分节符设置
            sectPr = OxmlElement('w:sectPr')
            
            # 设置分节符类型
            if section_type is not None and section_type != 'none':
                type_elem = OxmlElement('w:type')
                type_elem.set(qn('w:val'), section_type)
                sectPr.append(type_elem)
            
            # 添加基本页面设置
            # 页面大小
            pg_sz = OxmlElement('w:pgSz')
            pg_sz.set(qn('w:w'), '11906')
            pg_sz.set(qn('w:h'), '16838')
            sectPr.append(pg_sz)
            
            # 页边距
            pg_mar = OxmlElement('w:pgMar')
            pg_mar.set(qn('w:top'), '1440')
            pg_mar.set(qn('w:right'), '1800')
            pg_mar.set(qn('w:bottom'), '1440')
            pg_mar.set(qn('w:left'), '1800')
            pg_mar.set(qn('w:header'), '851')
            pg_mar.set(qn('w:footer'), '992')
            pg_mar.set(qn('w:gutter'), '0')
            sectPr.append(pg_mar)
            
            # 列设置
            cols = OxmlElement('w:cols')
            cols.set(qn('w:space'), '425')
            cols.set(qn('w:num'), '1')
            sectPr.append(cols)
            
            # 文档网格
            doc_grid = OxmlElement('w:docGrid')
            doc_grid.set(qn('w:type'), 'lines')
            doc_grid.set(qn('w:linePitch'), '312')
            doc_grid.set(qn('w:charSpace'), '0')
            sectPr.append(doc_grid)
        
        return sectPr
    
    def insert_section_break(self, after_paragraph_index, section_type='oddPage'):
        """
        在指定段落后插入分节符
        
        Args:
            after_paragraph_index: 段落索引（从0开始）
            section_type: 分节符类型 ('oddPage', 'evenPage', 'nextPage', 'continuous', 'none')
        """
        
        print(f"🎯 在第{after_paragraph_index + 1}段后插入{section_type}分节符")
        
        # 获取当前分节信息
        sections_info = self.get_current_sections_info()
        
        # 找到要插入分节符的位置属于哪个节
        target_section_info = None
        for section_info in sections_info:
            if section_info['paragraph_index'] > after_paragraph_index:
                target_section_info = section_info
                break
        
        if target_section_info is None:
            # 在最后一节中插入
            target_section_info = sections_info[-1] if sections_info else None
        
        # 获取目标段落
        if after_paragraph_index >= len(self.doc.paragraphs):
            print(f"❌ 段落索引超出范围")
            return False
        
        target_paragraph = self.doc.paragraphs[after_paragraph_index]
        
        # 创建新段落用于放置分节符
        new_p = OxmlElement('w:p')
        target_paragraph._p.addnext(new_p)
        
        # 创建分节符元素
        if target_section_info and 'sectPr' in target_section_info:
            # 复制现有分节符的页面设置
            sectPr = self.create_sectPr_element(
                section_type=section_type,
                copy_from_sectPr=target_section_info['sectPr']
            )
        else:
            # 创建新的分节符设置
            sectPr = self.create_sectPr_element(section_type=section_type)
        
        # 将分节符添加到新段落
        pPr = new_p.get_or_add_pPr()
        pPr.append(sectPr)
        
        print(f"✅ 成功插入{section_type}分节符")
        return True
    
    def clear_all_section_breaks(self):
        """清除所有段落内的分节符，只保留文档末尾body层的分节符"""
        print("🧹 开始清除所有分节符...")
        
        removed_count = 0
        paragraphs_to_remove = []
        
        # 遍历所有段落，查找并移除分节符
        for i, para in enumerate(self.doc.paragraphs):
            p_element = para._element
            pPr = p_element.pPr
            
            if pPr is not None:
                sectPr = pPr.find(qn('w:sectPr'))
                if sectPr is not None:
                    # 移除段落内的分节符
                    pPr.remove(sectPr)
                    removed_count += 1
                    
                    # 如果段落只有分节符，没有其他内容，标记为删除
                    para_text = para.text.strip()
                    if not para_text:
                        paragraphs_to_remove.append(i)
                    
                    print(f"  ✅ 移除第{i+1}段的分节符")
        
        # 删除空的分节符段落（从后往前删除，避免索引变化）
        for para_index in reversed(paragraphs_to_remove):
            para = self.doc.paragraphs[para_index]
            para._p.getparent().remove(para._p)
            print(f"  🗑️ 删除第{para_index+1}段（空分节符段落）")
        
        # 确保body层有分节符
        body_sectPr = self.body.find(qn('w:sectPr'))
        if body_sectPr is None:
            print("  ➕ 文档末尾没有分节符，创建默认分节符")
            default_sectPr = self.create_sectPr_element(section_type='none')
            self.body.append(default_sectPr)
        else:
            # 移除body层分节符的type属性（第一节不应该有type）
            body_sectType = body_sectPr.find(qn('w:type'))
            if body_sectType is not None:
                body_sectPr.remove(body_sectType)
                print("  🔧 移除body层分节符的type属性")
        
        print(f"✅ 清除完成，共移除{removed_count}个分节符")
        return removed_count
    
    def show_sections_structure(self):
        """显示文档分节结构"""
        print("📊 文档分节结构:")
        print("=" * 50)
        
        sections_info = self.get_current_sections_info()
        
        for i, section_info in enumerate(sections_info, 1):
            para_index = section_info['paragraph_index']
            section_type = section_info['type']
            
            if section_info.get('is_last_section'):
                print(f"第{i}节: 最后一节 (body末尾)")
            else:
                para_text = ""
                if para_index < len(self.doc.paragraphs):
                    para_text = self.doc.paragraphs[para_index].text[:30]
                    if len(para_text) > 30:
                        para_text += "..."
                
                print(f"第{i}节: 段落{para_index + 1} - {para_text}")
            
            if section_type:
                type_names = {
                    'oddPage': '奇数页分节符',
                    'evenPage': '偶数页分节符',
                    'nextPage': '分页符',
                    'continuous': '连续分节符'
                }
                print(f"      类型: {type_names.get(section_type, section_type)}")
            else:
                print(f"      类型: 无类型(第一节)")
            
            print()
    
    def save(self, output_path):
        """保存文档"""
        self.doc.save(output_path)
        print(f"💾 文档已保存: {output_path}")

def main():
    """主函数"""
    print("🎯 通用Word分节符处理器")
    print("=" * 60)
    
    # 初始化文档管理器
    doc_manager = WordSectionManager("test_document.docx")
    
    print("📋 修改前的分节结构:")
    doc_manager.show_sections_structure()
    
    print("\n" + "="*60)
    
    # 先清除所有分节符
    removed_count = doc_manager.clear_all_section_breaks()
    
    print(f"\n📋 清除后的分节结构:")
    doc_manager.show_sections_structure()
    
    print("\n" + "="*60)
    
    # 在第8段后插入奇数页分节符
    success = doc_manager.insert_section_break(
        after_paragraph_index=7,  # 第8段（索引从0开始）
        section_type='oddPage'
    )
    
    if success:
        print("\n📋 插入分节符后的结构:")
        doc_manager.show_sections_structure()
        
        # 保存文档
        doc_manager.save("clean_and_insert_section.docx")
    else:
        print("\n❌ 分节符插入失败!")

if __name__ == "__main__":
    main()