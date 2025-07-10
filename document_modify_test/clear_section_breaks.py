#!/usr/bin/env python3
"""
清除Word文档中所有分节符，只保留文档末尾body层的分节符
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def clear_all_section_breaks(file_path, output_path=None):
    """
    清除文档中所有段落内的分节符，只保留文档末尾body层的分节符
    
    Args:
        file_path: 输入文档路径
        output_path: 输出文档路径，如果为None则覆盖原文件
    
    Returns:
        dict: 包含操作结果的字典
    """
    
    print(f"🔄 开始处理文档: {file_path}")
    
    try:
        # 打开文档
        doc = Document(file_path)
        body = doc._body._element
        
        print(f"📄 文档共有 {len(doc.paragraphs)} 段")
        
        removed_count = 0
        empty_paragraphs_removed = 0
        paragraphs_to_remove = []
        
        print("🧹 开始清除段落内的分节符...")
        
        # 遍历所有段落，查找并移除分节符
        for i, para in enumerate(doc.paragraphs):
            p_element = para._element
            pPr = p_element.pPr
            
            if pPr is not None:
                sectPr = pPr.find(qn('w:sectPr'))
                if sectPr is not None:
                    # 检查段落是否为空（在移除分节符前检查）
                    para_text = para.text.strip()
                    
                    # 移除段落内的分节符
                    pPr.remove(sectPr)
                    removed_count += 1
                    
                    if not para_text:
                        # 空段落：分节符+段落一起删除
                        paragraphs_to_remove.append(i)
                        print(f"  ✅ 移除第{i+1}段的分节符（空段落，将一并删除）")
                    else:
                        # 非空段落：只移除分节符
                        print(f"  ✅ 移除第{i+1}段的分节符（保留段落内容）")
        
        # 删除空的分节符段落（从后往前删除，避免索引变化）
        print("🗑️ 删除空的分节符段落...")
        for para_index in reversed(paragraphs_to_remove):
            para = doc.paragraphs[para_index]
            para._p.getparent().remove(para._p)
            empty_paragraphs_removed += 1
            print(f"  🗑️ 删除第{para_index+1}段（空分节符段落）")
        
        # 处理文档末尾body层的分节符
        print("🔧 处理文档末尾的分节符...")
        body_sectPr = body.find(qn('w:sectPr'))
        
        if body_sectPr is None:
            print("  ➕ 文档末尾没有分节符，创建默认分节符")
            # 创建默认分节符（不带type属性）
            default_sectPr = create_default_sectPr()
            body.append(default_sectPr)
        else:
            print("  ✅ 文档末尾已有分节符")
            # 移除body层分节符的type属性（第一节不应该有type）
            body_sectType = body_sectPr.find(qn('w:type'))
            if body_sectType is not None:
                body_sectPr.remove(body_sectType)
                print("  🔧 移除body层分节符的type属性（第一节不应有type）")
        
        # 保存文档
        save_path = output_path if output_path else file_path
        doc.save(save_path)
        
        result = {
            'success': True,
            'removed_section_breaks': removed_count,
            'removed_empty_paragraphs': empty_paragraphs_removed,
            'total_paragraphs': len(doc.paragraphs),
            'output_path': save_path
        }
        
        print(f"✅ 清除完成!")
        print(f"   - 移除分节符: {removed_count}个")
        print(f"   - 删除空段落: {empty_paragraphs_removed}个")
        print(f"   - 剩余段落数: {len(doc.paragraphs)}个")
        print(f"💾 文档已保存: {save_path}")
        
        return result
        
    except Exception as e:
        print(f"❌ 处理失败: {e}")
        return {
            'success': False,
            'error': str(e)
        }

def create_default_sectPr():
    """创建默认的分节符属性（不带type属性）"""
    
    sectPr = OxmlElement('w:sectPr')
    
    # 页面大小 (A4)
    pg_sz = OxmlElement('w:pgSz')
    pg_sz.set(qn('w:w'), '11906')  # 8.27英寸 * 1440
    pg_sz.set(qn('w:h'), '16838')  # 11.69英寸 * 1440
    sectPr.append(pg_sz)
    
    # 页边距
    pg_mar = OxmlElement('w:pgMar')
    pg_mar.set(qn('w:top'), '1440')     # 1英寸
    pg_mar.set(qn('w:right'), '1800')   # 1.25英寸
    pg_mar.set(qn('w:bottom'), '1440')  # 1英寸
    pg_mar.set(qn('w:left'), '1800')    # 1.25英寸
    pg_mar.set(qn('w:header'), '851')   # 页眉距离
    pg_mar.set(qn('w:footer'), '992')   # 页脚距离
    pg_mar.set(qn('w:gutter'), '0')     # 装订线
    sectPr.append(pg_mar)
    
    # 列设置
    cols = OxmlElement('w:cols')
    cols.set(qn('w:space'), '425')  # 列间距
    cols.set(qn('w:num'), '1')      # 单列
    sectPr.append(cols)
    
    # 文档网格
    doc_grid = OxmlElement('w:docGrid')
    doc_grid.set(qn('w:type'), 'lines')
    doc_grid.set(qn('w:linePitch'), '312')
    doc_grid.set(qn('w:charSpace'), '0')
    sectPr.append(doc_grid)
    
    return sectPr

def show_document_sections(file_path):
    """显示文档的分节信息"""
    
    print(f"📊 查看文档分节信息: {file_path}")
    
    try:
        doc = Document(file_path)
        body = doc._body._element
        
        print(f"📄 文档共有 {len(doc.paragraphs)} 段, {len(doc.sections)} 节")
        print()
        
        section_count = 0
        
        # 检查段落内的分节符
        for i, para in enumerate(doc.paragraphs):
            p_element = para._element
            pPr = p_element.pPr
            
            if pPr is not None:
                sectPr = pPr.find(qn('w:sectPr'))
                if sectPr is not None:
                    section_count += 1
                    sectType = sectPr.find(qn('w:type'))
                    break_type = sectType.get(qn('w:val')) if sectType is not None else "无类型"
                    
                    para_text = para.text.strip()[:40]
                    if not para_text:
                        para_text = "(空段落)"
                    
                    print(f"第{section_count}节: 段落{i+1} - {para_text}")
                    print(f"       类型: {break_type}")
                    print()
        
        # 检查body层的分节符
        body_sectPr = body.find(qn('w:sectPr'))
        if body_sectPr is not None:
            section_count += 1
            body_sectType = body_sectPr.find(qn('w:type'))
            break_type = body_sectType.get(qn('w:val')) if body_sectType is not None else "无类型"
            
            print(f"第{section_count}节: 文档末尾 (body层)")
            print(f"       类型: {break_type}")
            print()
        
        print(f"📈 总计: {section_count}个分节")
        
    except Exception as e:
        print(f"❌ 查看失败: {e}")

def main():
    """主函数"""
    
    print("🎯 Word文档分节符清除工具")
    print("=" * 50)
    
    # 文件路径
    file_path = "test_document.docx"
    output_path = "cleaned_document.docx"
    
    # 显示清除前的分节信息
    print("📋 清除前的分节信息:")
    show_document_sections(file_path)
    
    print("\n" + "="*50)
    
    # 清除分节符
    result = clear_all_section_breaks(file_path, output_path)
    
    if result['success']:
        print("\n📋 清除后的分节信息:")
        show_document_sections(output_path)
    else:
        print(f"\n❌ 清除失败: {result['error']}")

if __name__ == "__main__":
    main()