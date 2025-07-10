from docx.oxml import OxmlElement
from docx.oxml.ns import qn  # 命名空间处理工具



def insert_section_after_paragraph(doc, target_paragraph_index, sect_properties):
    """
    在指定段落后插入包含分节符的新段落

    :param doc: Document对象
    :param target_paragraph_index: 目标段落索引（从0开始）
    :param sect_properties: 分节属性字典（示例格式见下方）
    """
    # 1. 获取目标段落对象并在其后插入新段落
    target_paragraph = doc.paragraphs[target_paragraph_index]
    
    # 创建新段落元素
    new_p = OxmlElement('w:p')
    
    # 在目标段落后插入
    target_paragraph._p.addnext(new_p)
    
    # 重新获取段落对象
    new_paragraph = doc.paragraphs[target_paragraph_index + 1]

    # 3. 创建分节属性元素 <w:sectPr>
    sect_pr = OxmlElement('w:sectPr')

    # 4. 设置分节符基础属性
    sect_pr.set(qn('w:rsidR'), sect_properties.get('rsidR', '0071767C'))
    sect_pr.set(qn('w:rsidRPr'), sect_properties.get('rsidRPr', '0071767C'))
    sect_pr.set(qn('w:rsidSect'), sect_properties.get('rsidSect', '0018700E'))

    # 5. 添加分节类型
    type_elem = OxmlElement('w:type')
    type_elem.set(qn('w:val'), sect_properties.get('type', 'evenPage'))
    sect_pr.append(type_elem)

    # 6. 添加页面尺寸
    pg_sz = OxmlElement('w:pgSz')
    pg_sz.set(qn('w:w'), str(sect_properties.get('pgSz_w', 11906)))
    pg_sz.set(qn('w:h'), str(sect_properties.get('pgSz_h', 16838)))
    sect_pr.append(pg_sz)

    # 7. 添加页边距
    pg_mar = OxmlElement('w:pgMar')
    margins = ['top', 'right', 'bottom', 'left', 'header', 'footer', 'gutter']
    for margin in margins:
        pg_mar.set(qn(f'w:{margin}'), str(sect_properties.get(f'pgMar_{margin}', 1440)))
    sect_pr.append(pg_mar)

    # 8. 添加列设置
    cols = OxmlElement('w:cols')
    cols.set(qn('w:space'), str(sect_properties.get('cols_space', 425)))
    sect_pr.append(cols)

    # 9. 添加文档网格
    doc_grid = OxmlElement('w:docGrid')
    doc_grid.set(qn('w:type'), sect_properties.get('docGrid_type', 'lines'))
    doc_grid.set(qn('w:linePitch'), str(sect_properties.get('docGrid_linePitch', 312)))
    sect_pr.append(doc_grid)

    # 10. 将分节符添加到新段落属性
    pPr = new_paragraph._p.get_or_add_pPr()
    pPr.append(sect_pr)


# 使用示例
if __name__ == "__main__":
    from docx import Document

    # 定义分节属性（实际使用时可外部传入）
    sect_config = {
        'rsidR': '0071767C',
        'rsidRPr': '0071767C',
        'rsidSect': '0018700E',
        'type': 'oddPage',
        'pgSz_w': 11906,
        'pgSz_h': 16838,
        'pgMar_top': 1440,
        'pgMar_right': 1800,
        'pgMar_bottom': 1440,
        'pgMar_left': 1800,
        'pgMar_header': 851,
        'pgMar_footer': 992,
        'pgMar_gutter': 0,
        'cols_space': 425,
        'docGrid_type': 'lines',
        'docGrid_linePitch': 312
    }

    doc = Document("test_document.docx")
    insert_section_after_paragraph(doc, target_paragraph_index=3, sect_properties=sect_config)
    doc.save("output.docx")