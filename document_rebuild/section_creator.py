"""
分节创建器 - 使用python-docx高级API创建分节、页码、页眉
"""

from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_toc_section(doc, toc_config, styles):
    """创建目录节（只创建内容，不设置页码）"""
    
    print("📋 创建目录节...")
    
    try:
        # 添加目录标题
        title_para = doc.add_paragraph()
        title_para.text = toc_config['title']
        title_para.style = doc.styles[styles['title']]
        
        # 添加目录字段（TOC）
        toc_para = doc.add_paragraph()
        run = toc_para.add_run()
        
        # 创建TOC字段
        fld_char1 = OxmlElement('w:fldChar')
        fld_char1.set(qn('w:fldCharType'), 'begin')
        
        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'  # 1-3级标题，超链接
        
        fld_char2 = OxmlElement('w:fldChar')
        fld_char2.set(qn('w:fldCharType'), 'separate')
        
        placeholder_text = OxmlElement('w:t')
        placeholder_text.text = "请在Word中右键更新目录"
        
        fld_char3 = OxmlElement('w:fldChar')
        fld_char3.set(qn('w:fldCharType'), 'end')
        
        # 添加到run
        run._r.append(fld_char1)
        run._r.append(instr_text)
        run._r.append(fld_char2)
        run._r.append(placeholder_text)
        run._r.append(fld_char3)
        
        print("  ✅ 目录内容创建完成")
        
        return True
        
    except Exception as e:
        print(f"❌ 创建目录节失败: {e}")
        return False

def create_chapter_section(doc, chapter_config, styles):
    """创建章节（只创建内容和分节，不设置页码）"""
    
    chapter_num = chapter_config['number']
    print(f"📚 创建第{chapter_num}章: {chapter_config['title']}")
    
    try:
        # 添加奇数页分节符（统一使用ODD_PAGE）
        section = doc.add_section(WD_SECTION_START.ODD_PAGE)
        
        # 设置分节属性
        setup_section_properties(section, chapter_config)
        
        # 添加章节内容
        add_chapter_content(doc, chapter_config, styles)
        
        print(f"  ✅ 第{chapter_num}章创建完成 ({len(chapter_config['content'])}段)")
        
        return section
        
    except Exception as e:
        print(f"❌ 创建第{chapter_num}章失败: {e}")
        return None

def setup_section_properties(section, chapter_config):
    """设置分节属性（使用高级API）"""
    try:
        # 可以在这里设置页面属性，如页边距、页面方向等
        # section.top_margin = Inches(1)
        # section.bottom_margin = Inches(1)
        # section.left_margin = Inches(1.25)
        # section.right_margin = Inches(1.25)
        
        return True
    except Exception as e:
        print(f"    设置分节属性失败: {e}")
        return False

def setup_page_numbers_for_all_sections(doc, structure):
    """统一设置所有分节的页码"""
    
    print("📄 设置页码...")
    
    try:
        sections = doc.sections
        
        # 1. 设置目录节页码（罗马数字）
        if len(sections) > 0:
            toc_section = sections[0]
            setup_page_numbers(toc_section, 'roman', 1, restart=True)
            print("  ✅ 目录节页码: 罗马数字从I开始")
        
        # 2. 设置章节页码（阿拉伯数字）
        for i, chapter in enumerate(structure['chapters']):
            section_idx = i + 1  # 目录是第0节，章节从第1节开始
            if section_idx < len(sections):
                chapter_section = sections[section_idx]
                
                if chapter['page_restart']:
                    # 第一章重新开始阿拉伯数字页码
                    setup_page_numbers(chapter_section, 'decimal', chapter['page_start'], restart=True)
                    print(f"  ✅ 第{chapter['number']}章页码: 阿拉伯数字从{chapter['page_start']}重新开始")
                else:
                    # 后续章节延续页码
                    setup_continuous_page_numbers(chapter_section)
                    print(f"  ✅ 第{chapter['number']}章页码: 阿拉伯数字延续编号")
        
        print(f"  📊 页码设置总结: 处理了{len(sections)}个分节")
        return True
        
    except Exception as e:
        print(f"❌ 设置页码失败: {e}")
        return False

def add_chapter_content(doc, chapter_config, styles):
    """添加章节内容"""
    try:
        for content_item in chapter_config['content']:
            para = doc.add_paragraph()
            para.text = content_item['text']
            
            # 应用样式
            content_type = content_item['type']
            if content_type in styles:
                para.style = doc.styles[styles[content_type]]
            
            print(f"    ✅ 添加{content_type}: {content_item['text'][:30]}...")
        
        return True
    except Exception as e:
        print(f"    添加章节内容失败: {e}")
        return False

def setup_page_numbers(section, format_type='decimal', start_number=None, restart=False):
    """设置页码（优化版 - 更好地使用API）"""
    
    try:
        # 断开与前节的链接
        footer = section.footer
        footer.is_linked_to_previous = False
        
        # 直接使用第一个段落，避免检查和添加逻辑
        footer_para = footer.paragraphs[0]
        footer_para.clear()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 根据格式类型添加页码
        if format_type == 'roman':
            # 罗马数字页码: I, II, III...
            page_run = footer_para.add_run()
            add_page_number_field(page_run)
            
            # 设置罗马数字格式（只在需要时）
            if start_number:
                set_page_number_format(section, 'upperRoman', start_number)
                
        else:
            # 阿拉伯数字页码: 第 1 页, 第 2 页...
            footer_para.add_run("第 ")
            page_run = footer_para.add_run()
            add_page_number_field(page_run)
            footer_para.add_run(" 页")
            
            # 只在需要重新开始时设置格式
            if restart and start_number:
                set_page_number_format(section, 'decimal', start_number)
        
        return True
        
    except Exception as e:
        print(f"    设置页码失败: {e}")
        return False

def set_page_number_format(section, format_name, start_number):
    """设置页码格式和起始编号"""
    try:
        sectPr = section._sectPr
        pgNumType = sectPr.find(qn('w:pgNumType'))
        if pgNumType is None:
            pgNumType = OxmlElement('w:pgNumType')
            sectPr.append(pgNumType)
        
        pgNumType.set(qn('w:fmt'), format_name)
        pgNumType.set(qn('w:start'), str(start_number))
        
        return True
    except Exception as e:
        print(f"    设置页码格式失败: {e}")
        return False

def add_page_number_field(run):
    """添加页码字段"""
    try:
        fld_char1 = OxmlElement('w:fldChar')
        fld_char1.set(qn('w:fldCharType'), 'begin')
        
        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        instr_text.text = "PAGE"
        
        fld_char2 = OxmlElement('w:fldChar')
        fld_char2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fld_char1)
        run._r.append(instr_text)
        run._r.append(fld_char2)
        
        return True
    except Exception as e:
        print(f"    添加页码字段失败: {e}")
        return False

def setup_continuous_page_numbers(section):
    """为后续章节设置连续页码（优化版 - 简化逻辑）"""
    
    try:
        # 断开与前节的链接  
        footer = section.footer
        footer.is_linked_to_previous = False
        
        # 直接使用第一个段落
        footer_para = footer.paragraphs[0]
        footer_para.clear()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 只添加页码字段，不设置任何格式XML
        # 页码会自动延续前面章节的编号
        footer_para.add_run("第 ")
        page_run = footer_para.add_run()
        add_page_number_field(page_run)
        footer_para.add_run(" 页")
        
        return True
        
    except Exception as e:
        print(f"    设置连续页码失败: {e}")
        return False

def setup_headers(doc, structure):
    """设置页眉（使用高级API）"""
    
    print("📄 设置页眉...")
    
    try:
        # 首先在文档级别启用奇偶页不同页眉页脚
        try:
            doc.settings.odd_and_even_pages_header_footer = True
            print("  ✅ 启用奇偶页不同页眉页脚")
        except:
            print("  ⚠️ 无法启用奇偶页不同（可能需要在Word中手动启用）")
        
        sections = doc.sections
        
        # 目录节页眉
        if len(sections) > 0:
            toc_section = sections[0]
            setup_section_headers(toc_section, structure['toc_section']['headers'])
            print("  ✅ 目录节页眉设置完成")
        
        # 章节页眉
        for i, chapter in enumerate(structure['chapters']):
            section_idx = i + 1  # 目录是第0节，章节从第1节开始
            if section_idx < len(sections):
                chapter_section = sections[section_idx]
                setup_section_headers(chapter_section, chapter['headers'])
                print(f"  ✅ 第{chapter['number']}章页眉设置完成")
        
        return True
        
    except Exception as e:
        print(f"❌ 设置页眉失败: {e}")
        return False

def setup_section_headers(section, header_config):
    """设置单个分节的页眉（使用优化的API）"""
    try:
        
        # 设置奇数页页眉（默认页眉）
        header = section.header
        header.is_linked_to_previous = False
        
        # 直接使用第一个段落，避免添加额外段落
        header_para = header.paragraphs[0]
        header_para.clear()
        header_para.text = header_config['odd']
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 设置偶数页页眉
        even_header = section.even_page_header  
        even_header.is_linked_to_previous = False
        
        # 直接使用第一个段落
        even_header_para = even_header.paragraphs[0]
        even_header_para.clear()
        even_header_para.text = header_config['even']
        even_header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        return True
        
    except Exception as e:
        print(f"    设置分节页眉失败: {e}")
        return False