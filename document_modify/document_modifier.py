#!/usr/bin/env python3
"""
主文档修改器 - 协调各个修改器基于配置修改Word文档
"""

import os
import time
from typing import Dict, Any, Optional
from docx import Document
from config_loader import ModifyConfigLoader, ModifyMode
from style_modifier import StyleModifier
from structure_modifier import StructureModifier

class DocumentModifier:
    """Word文档修改器主类"""
    
    def __init__(self, source_doc_path: str, config_path: str):
        """初始化文档修改器"""
        self.source_doc_path = source_doc_path
        self.config_path = config_path
        
        # 加载配置
        self.config_loader = ModifyConfigLoader(config_path)
        
        # 验证配置
        if not self.config_loader.validate_config():
            raise ValueError("配置文件验证失败")
        
        # 加载文档
        self.doc = Document(source_doc_path)
        
        # 初始化各个修改器
        self.style_modifier = StyleModifier(self.doc, self.config_loader)
        self.structure_modifier = StructureModifier(self.doc, self.config_loader)
        
        # 修改模式
        self.modify_mode = self.config_loader.get_modify_mode()
        
        print(f"✅ 文档修改器初始化完成")
        print(f"📄 源文档: {source_doc_path}")
        print(f"📋 配置文件: {config_path}")
        print(f"🔧 修改模式: {self.modify_mode.value}")
    
    def modify_document(self, output_path: Optional[str] = None) -> Dict[str, Any]:
        """
        执行文档修改
        
        Args:
            output_path: 输出文档路径，如果为None则覆盖原文档
            
        Returns:
            修改结果统计
        """
        
        print("\n" + "="*60)
        print("🚀 开始修改文档...")
        print("="*60)
        
        start_time = time.time()
        
        # 收集所有修改结果
        all_results = {
            'success': True,
            'modify_mode': self.modify_mode.value,
            'source_doc': self.source_doc_path,
            'output_doc': output_path or self.source_doc_path,
            'modifications': {},
            'errors': [],
            'warnings': [],
            'execution_time': 0
        }
        
        try:
            # 第一阶段：应用页面设置
            if self.config_loader.get_page_settings():
                print("\n📄 第一阶段：应用页面设置")
                page_results = self._apply_page_settings()
                all_results['modifications']['page_settings'] = page_results
            
            # 第二阶段：修改样式
            if self.config_loader.get_styles_config():
                print("\n🎨 第二阶段：修改样式")
                style_results = self.style_modifier.modify_all_styles()
                all_results['modifications']['styles'] = style_results
            
            # 第三阶段：修改文档结构
            if any([
                self.config_loader.get_structure_config(),
                self.config_loader.get_section_breaks_config(),
                self.config_loader.get_toc_settings(),
                self.config_loader.get_page_numbering_config(),
                self.config_loader.get_headers_footers_config()
            ]):
                print("\n📐 第三阶段：修改文档结构")
                structure_results = self.structure_modifier.modify_document_structure()
                all_results['modifications']['structure'] = structure_results
            
            # 第四阶段：应用内容修改
            content_mods = self.config_loader.get_content_modifications()
            if content_mods:
                print("\n📝 第四阶段：应用内容修改")
                content_results = self._apply_content_modifications(content_mods)
                all_results['modifications']['content'] = content_results
            
            # 第五阶段：最终处理和保存
            print("\n💾 第五阶段：保存文档")
            save_path = output_path or self.source_doc_path
            
            # 如果要覆盖原文件，先备份
            if not output_path and self.config_loader.get_document_info().get('backup_original', True):
                backup_path = self._create_backup()
                all_results['backup_path'] = backup_path
                print(f"  📦 已创建备份: {backup_path}")
            
            # 保存文档
            self.doc.save(save_path)
            print(f"  ✅ 文档已保存: {save_path}")
            
        except Exception as e:
            all_results['success'] = False
            all_results['errors'].append(str(e))
            print(f"\n❌ 修改过程出错: {e}")
        
        # 计算执行时间
        all_results['execution_time'] = round(time.time() - start_time, 2)
        
        # 显示结果摘要
        self._display_results_summary(all_results)
        
        return all_results
    
    def _apply_page_settings(self) -> Dict[str, Any]:
        """应用页面设置"""
        results = {
            'applied': False,
            'settings': {}
        }
        
        page_config = self.config_loader.get_page_settings()
        
        try:
            for section in self.doc.sections:
                # 页面大小
                if 'page_size' in page_config:
                    self._set_page_size(section, page_config['page_size'])
                    results['settings']['page_size'] = page_config['page_size']
                
                # 页面方向
                if 'orientation' in page_config:
                    orientation = self.config_loader.parse_orientation(page_config['orientation'])
                    if orientation:
                        section.orientation = orientation
                        results['settings']['orientation'] = page_config['orientation']
                
                # 页边距
                if 'margins' in page_config:
                    self._set_margins(section, page_config['margins'])
                    results['settings']['margins'] = page_config['margins']
                
                # 装订线
                if 'gutter' in page_config:
                    gutter = self.config_loader.parse_length(page_config['gutter'])
                    if gutter:
                        section.gutter = gutter
                        results['settings']['gutter'] = page_config['gutter']
            
            results['applied'] = True
            print("  ✅ 页面设置已应用")
            
        except Exception as e:
            results['error'] = str(e)
            print(f"  ❌ 页面设置失败: {e}")
        
        return results
    
    def _set_page_size(self, section: Any, page_size: str) -> None:
        """设置页面大小"""
        # 预定义页面大小
        page_sizes = {
            'A4': (210, 297),  # mm
            'A3': (297, 420),  # mm
            'Letter': (215.9, 279.4),  # mm
            'Legal': (215.9, 355.6),  # mm
        }
        
        if page_size in page_sizes:
            width_mm, height_mm = page_sizes[page_size]
            # 转换为EMU (English Metric Units)
            section.page_width = int(width_mm * 36000)
            section.page_height = int(height_mm * 36000)
    
    def _set_margins(self, section: Any, margins: Dict[str, str]) -> None:
        """设置页边距"""
        margin_names = ['top', 'bottom', 'left', 'right']
        
        for name in margin_names:
            if name in margins:
                margin_value = self.config_loader.parse_length(margins[name])
                if margin_value:
                    setattr(section, f'{name}_margin', margin_value)
    
    def _apply_content_modifications(self, content_mods: Dict[str, Any]) -> Dict[str, Any]:
        """应用内容修改"""
        results = {
            'replacements': 0,
            'insertions': 0,
            'deletions': 0,
            'formatting_changes': 0
        }
        
        # 文本替换
        if 'replacements' in content_mods:
            for replacement in content_mods['replacements']:
                count = self._replace_text(
                    replacement.get('find', ''),
                    replacement.get('replace', ''),
                    replacement.get('options', {})
                )
                results['replacements'] += count
        
        # 插入内容
        if 'insertions' in content_mods:
            for insertion in content_mods['insertions']:
                if self._insert_content(insertion):
                    results['insertions'] += 1
        
        # 删除内容
        if 'deletions' in content_mods:
            for deletion in content_mods['deletions']:
                if self._delete_content(deletion):
                    results['deletions'] += 1
        
        # 格式更改
        if 'formatting' in content_mods:
            for formatting in content_mods['formatting']:
                count = self._apply_formatting(formatting)
                results['formatting_changes'] += count
        
        return results
    
    def _replace_text(self, find_text: str, replace_text: str, options: Dict[str, Any]) -> int:
        """替换文本"""
        count = 0
        
        # 是否区分大小写
        case_sensitive = options.get('case_sensitive', True)
        # 是否全词匹配
        whole_word = options.get('whole_word', False)
        # 是否使用正则表达式
        use_regex = options.get('use_regex', False)
        
        # 遍历所有段落
        for para in self.doc.paragraphs:
            if find_text in para.text:
                if use_regex:
                    import re
                    para.text = re.sub(find_text, replace_text, para.text)
                else:
                    para.text = para.text.replace(find_text, replace_text)
                count += 1
        
        # 遍历表格
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if find_text in para.text:
                            para.text = para.text.replace(find_text, replace_text)
                            count += 1
        
        if count > 0:
            print(f"  ✅ 替换文本: '{find_text}' -> '{replace_text}' ({count}处)")
        
        return count
    
    def _insert_content(self, insertion: Dict[str, Any]) -> bool:
        """插入内容"""
        position = insertion.get('position', 'end')
        content = insertion.get('content', '')
        style = insertion.get('style', 'Normal')
        
        try:
            if position == 'beginning':
                para = self.doc.paragraphs[0].insert_paragraph_before(content)
            elif position == 'end':
                para = self.doc.add_paragraph(content)
            elif isinstance(position, int):
                if 0 <= position < len(self.doc.paragraphs):
                    para = self.doc.paragraphs[position].insert_paragraph_before(content)
                else:
                    para = self.doc.add_paragraph(content)
            else:
                return False
            
            # 应用样式
            if style in self.doc.styles:
                para.style = style
            
            print(f"  ✅ 插入内容: 位置={position}, 长度={len(content)}字符")
            return True
            
        except Exception as e:
            print(f"  ❌ 插入内容失败: {e}")
            return False
    
    def _delete_content(self, deletion: Dict[str, Any]) -> bool:
        """删除内容"""
        # 这个功能比较复杂，需要谨慎实现
        # 可以根据条件删除段落、表格等
        return False
    
    def _apply_formatting(self, formatting: Dict[str, Any]) -> int:
        """应用格式化"""
        count = 0
        criteria = formatting.get('criteria', {})
        format_settings = formatting.get('format', {})
        
        for para in self.doc.paragraphs:
            if self._match_formatting_criteria(para, criteria):
                self._apply_paragraph_formatting(para, format_settings)
                count += 1
        
        return count
    
    def _match_formatting_criteria(self, para: Any, criteria: Dict[str, Any]) -> bool:
        """匹配格式化条件"""
        # 实现条件匹配逻辑
        return True
    
    def _apply_paragraph_formatting(self, para: Any, settings: Dict[str, Any]) -> None:
        """应用段落格式化"""
        # 实现格式化逻辑
        pass
    
    def _create_backup(self) -> str:
        """创建备份文件"""
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        backup_dir = os.path.dirname(self.source_doc_path)
        backup_name = f"{os.path.splitext(os.path.basename(self.source_doc_path))[0]}_backup_{timestamp}.docx"
        backup_path = os.path.join(backup_dir, backup_name)
        
        # 复制原文件作为备份
        import shutil
        shutil.copy2(self.source_doc_path, backup_path)
        
        return backup_path
    
    def _display_results_summary(self, results: Dict[str, Any]) -> None:
        """显示结果摘要"""
        print("\n" + "="*60)
        print("📊 修改结果摘要")
        print("="*60)
        
        print(f"✅ 修改{'成功' if results['success'] else '失败'}")
        print(f"⏱️  执行时间: {results['execution_time']}秒")
        print(f"📄 输出文档: {results['output_doc']}")
        
        if 'backup_path' in results:
            print(f"📦 备份文件: {results['backup_path']}")
        
        # 显示各项修改统计
        mods = results.get('modifications', {})
        
        if 'styles' in mods:
            style_stats = mods['styles']
            print(f"\n🎨 样式修改:")
            print(f"  - 修改: {len(style_stats.get('modified_styles', []))}个")
            print(f"  - 创建: {len(style_stats.get('created_styles', []))}个")
            print(f"  - 映射: {len(style_stats.get('mapped_styles', []))}个")
        
        if 'structure' in mods:
            struct_stats = mods['structure']
            print(f"\n📐 结构修改:")
            print(f"  - 分节符: {struct_stats.get('sections_modified', 0)}个")
            print(f"  - 标题更新: {struct_stats.get('headings_updated', 0)}个")
            print(f"  - 页码更新: {'是' if struct_stats.get('page_numbers_updated') else '否'}")
            print(f"  - 页眉页脚更新: {'是' if struct_stats.get('headers_footers_updated') else '否'}")
        
        if 'content' in mods:
            content_stats = mods['content']
            print(f"\n📝 内容修改:")
            print(f"  - 文本替换: {content_stats.get('replacements', 0)}处")
            print(f"  - 内容插入: {content_stats.get('insertions', 0)}处")
            print(f"  - 内容删除: {content_stats.get('deletions', 0)}处")
            print(f"  - 格式更改: {content_stats.get('formatting_changes', 0)}处")
        
        # 显示错误和警告
        if results.get('errors'):
            print(f"\n❌ 错误 ({len(results['errors'])}个):")
            for error in results['errors']:
                print(f"  - {error}")
        
        if results.get('warnings'):
            print(f"\n⚠️  警告 ({len(results['warnings'])}个):")
            for warning in results['warnings']:
                print(f"  - {warning}")
        
        print("\n" + "="*60)

def main():
    """测试主文档修改器"""
    
    # 创建测试文档
    test_doc = Document()
    test_doc.add_heading('测试文档', 0)
    test_doc.add_heading('第一章 介绍', 1)
    test_doc.add_paragraph('这是一个测试文档，用于演示文档修改器的功能。')
    test_doc.add_heading('第二章 内容', 1)
    test_doc.add_paragraph('这里是第二章的内容。')
    test_doc.add_paragraph('需要被替换的文本：旧文本')
    
    test_doc_path = "test_source_document.docx"
    test_doc.save(test_doc_path)
    
    # 创建测试配置
    import json
    
    test_config = {
        "modify_mode": "merge",
        "document_info": {
            "backup_original": True
        },
        "page_settings": {
            "page_size": "A4",
            "orientation": "portrait",
            "margins": {
                "top": "2.54cm",
                "bottom": "2.54cm",
                "left": "3.18cm",
                "right": "3.18cm"
            }
        },
        "styles": {
            "Heading 1": {
                "font": {
                    "name": "黑体",
                    "size": "16pt",
                    "bold": True,
                    "color": "#000080"
                },
                "paragraph": {
                    "alignment": "left",
                    "space_before": "12pt",
                    "space_after": "6pt"
                }
            }
        },
        "content_modifications": {
            "replacements": [
                {
                    "find": "旧文本",
                    "replace": "新文本",
                    "options": {
                        "case_sensitive": True
                    }
                }
            ],
            "insertions": [
                {
                    "position": "end",
                    "content": "这是新增加的段落。",
                    "style": "Normal"
                }
            ]
        }
    }
    
    # 保存配置
    config_path = "test_modify_config.json"
    with open(config_path, 'w', encoding='utf-8') as f:
        json.dump(test_config, f, ensure_ascii=False, indent=2)
    
    # 测试修改器
    print("🎯 测试文档修改器")
    print("="*60)
    
    modifier = DocumentModifier(test_doc_path, config_path)
    results = modifier.modify_document("test_modified_document.docx")
    
    print(f"\n✅ 测试完成")
    
    # 清理测试文件
    import os
    os.remove(test_doc_path)
    os.remove(config_path)
    if os.path.exists("test_modified_document.docx"):
        os.remove("test_modified_document.docx")
    
    # 清理备份文件
    for file in os.listdir('.'):
        if file.startswith('test_source_document_backup_'):
            os.remove(file)

if __name__ == "__main__":
    main()