#!/usr/bin/env python3
"""
样式修改器 - 基于配置修改Word文档样式
"""

from typing import Dict, Any, Optional, List
from docx import Document
from docx.styles.style import _ParagraphStyle, _CharacterStyle, _TableStyle
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from config_loader import ModifyConfigLoader, ModifyMode

class StyleModifier:
    """Word文档样式修改器"""
    
    def __init__(self, document: Document, config_loader: ModifyConfigLoader):
        """初始化样式修改器"""
        self.doc = document
        self.config_loader = config_loader
        self.styles_config = config_loader.get_styles_config()
        self.modify_mode = config_loader.get_modify_mode()
        self.style_mapping = config_loader.get_style_mapping()
        
    def modify_all_styles(self) -> Dict[str, Any]:
        """修改所有样式"""
        results = {
            'modified_styles': [],
            'created_styles': [],
            'mapped_styles': [],
            'errors': []
        }
        
        print("🎨 开始修改文档样式...")
        
        # 处理样式配置
        for style_name, style_config in self.styles_config.items():
            try:
                # 检查是否需要映射样式名
                target_style_name = self.style_mapping.get(style_name, style_name)
                
                # 尝试获取现有样式
                existing_style = self._get_style_by_name(target_style_name)
                
                if existing_style:
                    # 修改现有样式
                    self._modify_existing_style(existing_style, style_config)
                    results['modified_styles'].append(target_style_name)
                    print(f"  ✅ 修改样式: {target_style_name}")
                else:
                    # 创建新样式
                    if self.modify_mode in [ModifyMode.MERGE, ModifyMode.APPEND]:
                        new_style = self._create_new_style(target_style_name, style_config)
                        if new_style:
                            results['created_styles'].append(target_style_name)
                            print(f"  ➕ 创建样式: {target_style_name}")
                    else:
                        print(f"  ⚠️ 样式不存在且模式不允许创建: {target_style_name}")
                
                # 记录映射
                if style_name != target_style_name:
                    results['mapped_styles'].append(f"{style_name} -> {target_style_name}")
                    
            except Exception as e:
                error_msg = f"处理样式 {style_name} 时出错: {str(e)}"
                results['errors'].append(error_msg)
                print(f"  ❌ {error_msg}")
        
        # 应用批量样式操作
        self._apply_batch_style_operations()
        
        return results
    
    def _get_style_by_name(self, style_name: str) -> Optional[Any]:
        """根据名称获取样式"""
        try:
            return self.doc.styles[style_name]
        except KeyError:
            return None
    
    def _modify_existing_style(self, style: Any, config: Dict[str, Any]) -> None:
        """修改现有样式"""
        
        # 根据修改模式处理
        if self.modify_mode == ModifyMode.REPLACE:
            # 替换模式：先重置样式
            self._reset_style(style)
        
        # 应用字体设置
        if 'font' in config and hasattr(style, 'font'):
            self._apply_font_settings(style.font, config['font'])
        
        # 应用段落设置
        if 'paragraph' in config and hasattr(style, 'paragraph_format'):
            self._apply_paragraph_settings(style.paragraph_format, config['paragraph'])
        
        # 应用其他设置
        if 'base_style' in config and hasattr(style, 'base_style'):
            base_style = self._get_style_by_name(config['base_style'])
            if base_style:
                style.base_style = base_style
        
        if 'next_paragraph_style' in config and hasattr(style, 'next_paragraph_style'):
            next_style = self._get_style_by_name(config['next_paragraph_style'])
            if next_style:
                style.next_paragraph_style = next_style
    
    def _create_new_style(self, style_name: str, config: Dict[str, Any]) -> Optional[Any]:
        """创建新样式"""
        
        # 确定样式类型
        style_type = self._determine_style_type(config)
        
        try:
            # 创建样式
            new_style = self.doc.styles.add_style(style_name, style_type)
            
            # 应用配置
            self._modify_existing_style(new_style, config)
            
            return new_style
            
        except Exception as e:
            print(f"    创建样式失败: {e}")
            return None
    
    def _determine_style_type(self, config: Dict[str, Any]) -> WD_STYLE_TYPE:
        """根据配置确定样式类型"""
        
        # 显式指定的类型
        if 'type' in config:
            type_map = {
                'paragraph': WD_STYLE_TYPE.PARAGRAPH,
                'character': WD_STYLE_TYPE.CHARACTER,
                'table': WD_STYLE_TYPE.TABLE,
                'list': WD_STYLE_TYPE.LIST
            }
            return type_map.get(config['type'].lower(), WD_STYLE_TYPE.PARAGRAPH)
        
        # 根据配置内容推断
        if 'paragraph' in config:
            return WD_STYLE_TYPE.PARAGRAPH
        elif 'table' in config:
            return WD_STYLE_TYPE.TABLE
        else:
            return WD_STYLE_TYPE.CHARACTER
    
    def _reset_style(self, style: Any) -> None:
        """重置样式到默认状态"""
        # 这里可以实现样式重置逻辑
        pass
    
    def _apply_font_settings(self, font: Any, config: Dict[str, Any]) -> None:
        """应用字体设置"""
        
        # 字体名称
        if 'name' in config:
            font.name = config['name']
        
        # 字体大小
        if 'size' in config:
            size = self.config_loader.parse_length(config['size'])
            if size:
                font.size = size
        
        # 粗体
        if 'bold' in config:
            font.bold = self.config_loader.parse_boolean(config['bold'])
        
        # 斜体
        if 'italic' in config:
            font.italic = self.config_loader.parse_boolean(config['italic'])
        
        # 下划线
        if 'underline' in config:
            font.underline = self.config_loader.parse_boolean(config['underline'])
        
        # 删除线
        if 'strike' in config:
            font.strike = self.config_loader.parse_boolean(config['strike'])
        
        # 颜色
        if 'color' in config:
            color_tuple = self.config_loader.parse_color(config['color'])
            if color_tuple:
                font.color.rgb = RGBColor(*color_tuple)
        
        # 主题颜色
        if 'theme_color' in config:
            font.color.theme_color = self._parse_theme_color(config['theme_color'])
        
        # 其他效果
        if 'all_caps' in config:
            font.all_caps = self.config_loader.parse_boolean(config['all_caps'])
        
        if 'small_caps' in config:
            font.small_caps = self.config_loader.parse_boolean(config['small_caps'])
        
        if 'shadow' in config:
            font.shadow = self.config_loader.parse_boolean(config['shadow'])
        
        if 'outline' in config:
            font.outline = self.config_loader.parse_boolean(config['outline'])
        
        if 'emboss' in config:
            font.emboss = self.config_loader.parse_boolean(config['emboss'])
        
        if 'imprint' in config:
            font.imprint = self.config_loader.parse_boolean(config['imprint'])
        
        # 上下标
        if 'superscript' in config:
            font.superscript = self.config_loader.parse_boolean(config['superscript'])
        
        if 'subscript' in config:
            font.subscript = self.config_loader.parse_boolean(config['subscript'])
    
    def _apply_paragraph_settings(self, paragraph_format: Any, config: Dict[str, Any]) -> None:
        """应用段落设置"""
        
        # 对齐方式
        if 'alignment' in config:
            alignment = self.config_loader.parse_alignment(config['alignment'])
            if alignment:
                paragraph_format.alignment = alignment
        
        # 缩进
        if 'left_indent' in config:
            indent = self.config_loader.parse_length(config['left_indent'])
            if indent:
                paragraph_format.left_indent = indent
        
        if 'right_indent' in config:
            indent = self.config_loader.parse_length(config['right_indent'])
            if indent:
                paragraph_format.right_indent = indent
        
        if 'first_line_indent' in config:
            indent = self.config_loader.parse_length(config['first_line_indent'])
            if indent:
                paragraph_format.first_line_indent = indent
        
        # 间距
        if 'space_before' in config:
            space = self.config_loader.parse_length(config['space_before'])
            if space:
                paragraph_format.space_before = space
        
        if 'space_after' in config:
            space = self.config_loader.parse_length(config['space_after'])
            if space:
                paragraph_format.space_after = space
        
        # 行距
        if 'line_spacing' in config:
            paragraph_format.line_spacing = float(config['line_spacing'])
        
        if 'line_spacing_rule' in config:
            rule = self.config_loader.parse_line_spacing_rule(config['line_spacing_rule'])
            if rule:
                paragraph_format.line_spacing_rule = rule
        
        # 分页控制
        if 'keep_together' in config:
            paragraph_format.keep_together = self.config_loader.parse_boolean(config['keep_together'])
        
        if 'keep_with_next' in config:
            paragraph_format.keep_with_next = self.config_loader.parse_boolean(config['keep_with_next'])
        
        if 'page_break_before' in config:
            paragraph_format.page_break_before = self.config_loader.parse_boolean(config['page_break_before'])
        
        if 'widow_control' in config:
            paragraph_format.widow_control = self.config_loader.parse_boolean(config['widow_control'])
    
    def _parse_theme_color(self, theme_color: str) -> Optional[MSO_THEME_COLOR_INDEX]:
        """解析主题颜色"""
        theme_color_map = {
            'accent_1': MSO_THEME_COLOR_INDEX.ACCENT_1,
            'accent_2': MSO_THEME_COLOR_INDEX.ACCENT_2,
            'accent_3': MSO_THEME_COLOR_INDEX.ACCENT_3,
            'accent_4': MSO_THEME_COLOR_INDEX.ACCENT_4,
            'accent_5': MSO_THEME_COLOR_INDEX.ACCENT_5,
            'accent_6': MSO_THEME_COLOR_INDEX.ACCENT_6,
            'background_1': MSO_THEME_COLOR_INDEX.BACKGROUND_1,
            'background_2': MSO_THEME_COLOR_INDEX.BACKGROUND_2,
            'dark_1': MSO_THEME_COLOR_INDEX.DARK_1,
            'dark_2': MSO_THEME_COLOR_INDEX.DARK_2,
            'light_1': MSO_THEME_COLOR_INDEX.LIGHT_1,
            'light_2': MSO_THEME_COLOR_INDEX.LIGHT_2,
            'hyperlink': MSO_THEME_COLOR_INDEX.HYPERLINK,
            'followed_hyperlink': MSO_THEME_COLOR_INDEX.FOLLOWED_HYPERLINK
        }
        return theme_color_map.get(theme_color.lower())
    
    def _apply_batch_style_operations(self) -> None:
        """应用批量样式操作"""
        batch_ops = self.config_loader.get_batch_operations()
        
        for op in batch_ops:
            op_type = op.get('type')
            
            if op_type == 'apply_style_to_all':
                # 将指定样式应用到所有符合条件的段落
                self._apply_style_to_paragraphs(op)
            elif op_type == 'replace_style':
                # 替换文档中的样式使用
                self._replace_style_usage(op)
            elif op_type == 'remove_unused_styles':
                # 移除未使用的样式
                self._remove_unused_styles()
    
    def _apply_style_to_paragraphs(self, operation: Dict[str, Any]) -> None:
        """将样式应用到段落"""
        style_name = operation.get('style_name')
        condition = operation.get('condition', {})
        
        style = self._get_style_by_name(style_name)
        if not style:
            return
        
        for para in self.doc.paragraphs:
            if self._match_paragraph_condition(para, condition):
                para.style = style
    
    def _match_paragraph_condition(self, paragraph: Any, condition: Dict[str, Any]) -> bool:
        """检查段落是否符合条件"""
        # 这里可以实现复杂的条件匹配逻辑
        if 'text_contains' in condition:
            return condition['text_contains'] in paragraph.text
        if 'style_name' in condition:
            return paragraph.style.name == condition['style_name']
        return True
    
    def _replace_style_usage(self, operation: Dict[str, Any]) -> None:
        """替换样式使用"""
        old_style_name = operation.get('old_style')
        new_style_name = operation.get('new_style')
        
        old_style = self._get_style_by_name(old_style_name)
        new_style = self._get_style_by_name(new_style_name)
        
        if not old_style or not new_style:
            return
        
        # 替换段落样式
        for para in self.doc.paragraphs:
            if para.style == old_style:
                para.style = new_style
        
        # 替换运行样式
        for para in self.doc.paragraphs:
            for run in para.runs:
                if hasattr(run, 'style') and run.style == old_style:
                    run.style = new_style
    
    def _remove_unused_styles(self) -> None:
        """移除未使用的自定义样式"""
        # 收集使用中的样式
        used_styles = set()
        
        # 检查段落样式
        for para in self.doc.paragraphs:
            if para.style:
                used_styles.add(para.style.name)
        
        # 检查运行样式
        for para in self.doc.paragraphs:
            for run in para.runs:
                if hasattr(run, 'style') and run.style:
                    used_styles.add(run.style.name)
        
        # 检查表格样式
        for table in self.doc.tables:
            if hasattr(table, 'style') and table.style:
                used_styles.add(table.style.name)
        
        # 注意：python-docx不支持直接删除样式
        # 这里只是标记未使用的样式
        for style in self.doc.styles:
            if not style.builtin and style.name not in used_styles:
                print(f"  🗑️ 未使用的样式: {style.name}")
    
    def get_all_styles_info(self) -> List[Dict[str, Any]]:
        """获取所有样式信息"""
        styles_info = []
        
        for style in self.doc.styles:
            info = {
                'name': style.name,
                'type': str(style.type),
                'builtin': style.builtin,
                'base_style': style.base_style.name if hasattr(style, 'base_style') and style.base_style else None
            }
            styles_info.append(info)
        
        return styles_info

def main():
    """测试样式修改器"""
    
    # 创建测试文档
    doc = Document()
    doc.add_heading('测试标题', 1)
    doc.add_paragraph('测试段落')
    
    # 创建测试配置
    from config_loader import ModifyConfigLoader
    
    config = {
        "modify_mode": "merge",
        "styles": {
            "Heading 1": {
                "font": {
                    "name": "微软雅黑",
                    "size": "18pt",
                    "bold": true,
                    "color": "#0000FF"
                },
                "paragraph": {
                    "alignment": "center",
                    "space_before": "12pt",
                    "space_after": "6pt"
                }
            },
            "Normal": {
                "font": {
                    "name": "宋体",
                    "size": "12pt"
                },
                "paragraph": {
                    "line_spacing": 1.5,
                    "first_line_indent": "2em"
                }
            }
        }
    }
    
    # 保存配置
    import json
    config_path = "test_style_config.json"
    with open(config_path, 'w', encoding='utf-8') as f:
        json.dump(config, f, ensure_ascii=False, indent=2)
    
    # 测试修改器
    loader = ModifyConfigLoader(config_path)
    modifier = StyleModifier(doc, loader)
    
    results = modifier.modify_all_styles()
    
    print(f"\n修改结果:")
    print(f"修改的样式: {results['modified_styles']}")
    print(f"创建的样式: {results['created_styles']}")
    print(f"错误: {results['errors']}")
    
    # 保存文档
    doc.save("test_styled_document.docx")
    print(f"\n✅ 测试文档已保存")
    
    # 清理
    import os
    os.remove(config_path)
    os.remove("test_styled_document.docx")

if __name__ == "__main__":
    main()