#!/usr/bin/env python3
"""
结构修改器 - 基于配置修改Word文档结构
"""

from typing import Dict, Any, List, Optional, Tuple
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from config_loader import ModifyConfigLoader, ModifyMode
import re

class StructureModifier:
    """Word文档结构修改器"""
    
    def __init__(self, document: Document, config_loader: ModifyConfigLoader):
        """初始化结构修改器"""
        self.doc = document
        self.config_loader = config_loader
        self.structure_config = config_loader.get_structure_config()
        self.modify_mode = config_loader.get_modify_mode()
        self.section_breaks_config = config_loader.get_section_breaks_config()
        self.toc_settings = config_loader.get_toc_settings()
        self.selective_mods = config_loader.get_selective_modifications()
        
    def modify_document_structure(self) -> Dict[str, Any]:
        """修改文档结构"""
        results = {
            'sections_modified': 0,
            'toc_operations': [],
            'headings_updated': 0,
            'page_numbers_updated': False,
            'headers_footers_updated': False,
            'errors': []
        }
        
        print("📐 开始修改文档结构...")
        
        try:
            # 1. 修改分节符
            if self.section_breaks_config:
                sections_result = self._modify_section_breaks()
                results['sections_modified'] = sections_result
            
            # 2. 修改或添加目录
            if self.toc_settings:
                toc_result = self._modify_toc()
                results['toc_operations'] = toc_result
            
            # 3. 修改标题层级
            if 'heading_levels' in self.structure_config:
                headings_result = self._modify_heading_levels()
                results['headings_updated'] = headings_result
            
            # 4. 修改页码
            page_num_config = self.config_loader.get_page_numbering_config()
            if page_num_config:
                results['page_numbers_updated'] = self._modify_page_numbering(page_num_config)
            
            # 5. 修改页眉页脚
            headers_footers_config = self.config_loader.get_headers_footers_config()
            if headers_footers_config:
                results['headers_footers_updated'] = self._modify_headers_footers(headers_footers_config)
            
            # 6. 应用选择性修改
            if self.selective_mods:
                self._apply_selective_modifications()
                
        except Exception as e:
            results['errors'].append(f"结构修改错误: {str(e)}")
            print(f"  ❌ {str(e)}")
        
        return results
    
    def _modify_section_breaks(self) -> int:
        """修改分节符"""
        modified_count = 0
        
        if self.modify_mode == ModifyMode.REPLACE:
            # 替换模式：先清除所有分节符
            self._clear_all_section_breaks()
        
        # 处理分节符配置
        for section_config in self.section_breaks_config.get('sections', []):
            after_paragraph = section_config.get('after_paragraph')
            break_type = section_config.get('type', 'new_page')
            
            if after_paragraph is not None:
                if self._add_section_break(after_paragraph, break_type):
                    modified_count += 1
                    print(f"  ✅ 在第{after_paragraph + 1}段后添加{break_type}分节符")
        
        return modified_count
    
    def _clear_all_section_breaks(self) -> None:
        """清除所有分节符（保留最后一个）"""
        # 遍历所有段落，移除分节符
        for para in self.doc.paragraphs:
            p_element = para._element
            pPr = p_element.pPr
            if pPr is not None:
                sectPr = pPr.find(qn('w:sectPr'))
                if sectPr is not None:
                    pPr.remove(sectPr)
        
        print("  🧹 清除所有段落内的分节符")
    
    def _add_section_break(self, after_paragraph_index: int, break_type: str) -> bool:
        """在指定段落后添加分节符"""
        try:
            if after_paragraph_index >= len(self.doc.paragraphs):
                return False
            
            # 获取目标段落
            target_paragraph = self.doc.paragraphs[after_paragraph_index]
            
            # 创建新段落用于放置分节符
            new_p = OxmlElement('w:p')
            target_paragraph._p.addnext(new_p)
            
            # 创建分节符
            sectPr = self._create_section_properties(break_type)
            
            # 将分节符添加到新段落
            pPr = new_p.get_or_add_pPr()
            pPr.append(sectPr)
            
            return True
            
        except Exception as e:
            print(f"    添加分节符失败: {e}")
            return False
    
    def _create_section_properties(self, break_type: str) -> OxmlElement:
        """创建分节符属性"""
        sectPr = OxmlElement('w:sectPr')
        
        # 设置分节符类型
        parsed_type = self.config_loader.parse_section_break_type(break_type)
        if parsed_type:
            type_elem = OxmlElement('w:type')
            type_elem.set(qn('w:val'), break_type)
            sectPr.append(type_elem)
        
        # 添加页面设置（从配置获取）
        page_settings = self.config_loader.get_page_settings()
        
        # 页面大小
        pg_sz = OxmlElement('w:pgSz')
        pg_sz.set(qn('w:w'), '11906')  # A4宽度
        pg_sz.set(qn('w:h'), '16838')  # A4高度
        sectPr.append(pg_sz)
        
        # 页边距
        pg_mar = OxmlElement('w:pgMar')
        margins = page_settings.get('margins', {})
        pg_mar.set(qn('w:top'), str(self._parse_margin(margins.get('top', '2.54cm'))))
        pg_mar.set(qn('w:right'), str(self._parse_margin(margins.get('right', '3.18cm'))))
        pg_mar.set(qn('w:bottom'), str(self._parse_margin(margins.get('bottom', '2.54cm'))))
        pg_mar.set(qn('w:left'), str(self._parse_margin(margins.get('left', '3.18cm'))))
        pg_mar.set(qn('w:header'), '851')
        pg_mar.set(qn('w:footer'), '992')
        pg_mar.set(qn('w:gutter'), '0')
        sectPr.append(pg_mar)
        
        return sectPr
    
    def _parse_margin(self, margin_str: str) -> int:
        """解析边距值为twips"""
        length = self.config_loader.parse_length(margin_str)
        if length:
            # 转换为twips (1/20 pt)
            return int(length.pt * 20)
        return 1440  # 默认1英寸
    
    def _modify_toc(self) -> List[str]:
        """修改或添加目录"""
        operations = []
        
        toc_config = self.toc_settings
        operation = toc_config.get('operation', 'update')  # update, add, remove
        
        if operation == 'add':
            # 添加目录
            position = toc_config.get('position', 'after_title')
            if self._add_toc(position, toc_config):
                operations.append("添加目录")
                print("  ✅ 添加目录")
        
        elif operation == 'update':
            # 更新现有目录
            if self._update_toc(toc_config):
                operations.append("更新目录")
                print("  ✅ 更新目录")
        
        elif operation == 'remove':
            # 移除目录
            if self._remove_toc():
                operations.append("移除目录")
                print("  ✅ 移除目录")
        
        return operations
    
    def _add_toc(self, position: str, config: Dict[str, Any]) -> bool:
        """添加目录"""
        try:
            # 查找插入位置
            insert_index = self._find_toc_insert_position(position)
            
            # 插入目录标题
            toc_title = config.get('title', '目录')
            toc_title_style = config.get('title_style', 'TOC Heading')
            
            # 在指定位置插入段落
            if insert_index >= len(self.doc.paragraphs):
                toc_para = self.doc.add_paragraph(toc_title)
            else:
                # 在现有段落之前插入
                toc_para = self.doc.paragraphs[insert_index].insert_paragraph_before(toc_title)
            
            # 应用样式
            if toc_title_style in self.doc.styles:
                toc_para.style = toc_title_style
            
            # 添加目录字段代码（简化版）
            toc_para = toc_para._p.addnext(self._create_toc_field(config))
            
            return True
            
        except Exception as e:
            print(f"    添加目录失败: {e}")
            return False
    
    def _create_toc_field(self, config: Dict[str, Any]) -> OxmlElement:
        """创建目录字段"""
        # 这是一个简化的TOC字段创建
        # 实际的TOC字段非常复杂，这里只是示例
        
        p = OxmlElement('w:p')
        pPr = p.get_or_add_pPr()
        
        # 添加目录样式
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), 'TOC1')
        pPr.append(pStyle)
        
        # 添加示例文本
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = "[目录将在此处生成]"
        r.append(t)
        p.append(r)
        
        return p
    
    def _find_toc_insert_position(self, position: str) -> int:
        """查找目录插入位置"""
        if position == 'after_title':
            # 在第一个标题之后
            for i, para in enumerate(self.doc.paragraphs):
                if para.style.name.startswith('Title') or para.style.name.startswith('Heading'):
                    return i + 1
            return 0
        elif position == 'beginning':
            return 0
        elif isinstance(position, int):
            return position
        else:
            return 1
    
    def _update_toc(self, config: Dict[str, Any]) -> bool:
        """更新目录"""
        # python-docx不支持直接更新TOC字段
        # 这里只是标记需要更新
        print("    ⚠️ 目录需要在Word中手动更新（右键->更新字段）")
        return True
    
    def _remove_toc(self) -> bool:
        """移除目录"""
        # 查找并移除目录相关段落
        # 这需要识别TOC字段，比较复杂
        return False
    
    def _modify_heading_levels(self) -> int:
        """修改标题层级"""
        modified_count = 0
        heading_config = self.structure_config.get('heading_levels', {})
        
        for para in self.doc.paragraphs:
            if para.style.name.startswith('Heading'):
                # 获取当前级别
                current_level = self._get_heading_level(para.style.name)
                
                # 检查是否需要修改
                for rule in heading_config.get('rules', []):
                    if self._match_heading_rule(para, current_level, rule):
                        new_level = rule.get('new_level')
                        if new_level and new_level != current_level:
                            para.style = f'Heading {new_level}'
                            modified_count += 1
                            print(f"  ✅ 修改标题级别: {current_level} -> {new_level}")
        
        return modified_count
    
    def _get_heading_level(self, style_name: str) -> Optional[int]:
        """获取标题级别"""
        match = re.match(r'Heading (\d+)', style_name)
        if match:
            return int(match.group(1))
        return None
    
    def _match_heading_rule(self, para: Any, level: int, rule: Dict[str, Any]) -> bool:
        """匹配标题规则"""
        # 级别匹配
        if 'from_level' in rule and rule['from_level'] != level:
            return False
        
        # 文本匹配
        if 'text_pattern' in rule:
            pattern = rule['text_pattern']
            if not re.search(pattern, para.text):
                return False
        
        # 位置匹配
        if 'position' in rule:
            # 这里可以实现更复杂的位置匹配逻辑
            pass
        
        return True
    
    def _modify_page_numbering(self, config: Dict[str, Any]) -> bool:
        """修改页码设置"""
        try:
            for section in self.doc.sections:
                # 设置页码格式
                if 'format' in config:
                    self._set_page_number_format(section, config['format'])
                
                # 设置起始页码
                if 'start_number' in config:
                    section.start_type = 0  # 继续前一节
                    section.page_number_start = config['start_number']
                
                # 设置页码位置
                if 'position' in config:
                    self._set_page_number_position(section, config['position'])
            
            print("  ✅ 更新页码设置")
            return True
            
        except Exception as e:
            print(f"  ❌ 页码设置失败: {e}")
            return False
    
    def _set_page_number_format(self, section: Any, format_type: str) -> None:
        """设置页码格式"""
        # 这需要操作底层XML
        # python-docx对页码的支持有限
        pass
    
    def _set_page_number_position(self, section: Any, position: str) -> None:
        """设置页码位置"""
        # 这通常在页眉或页脚中设置
        pass
    
    def _modify_headers_footers(self, config: Dict[str, Any]) -> bool:
        """修改页眉页脚"""
        try:
            for section in self.doc.sections:
                # 处理页眉
                if 'header' in config:
                    self._modify_header(section.header, config['header'])
                
                # 处理页脚
                if 'footer' in config:
                    self._modify_footer(section.footer, config['footer'])
                
                # 处理首页不同
                if 'different_first_page' in config:
                    section.different_first_page_header_footer = config['different_first_page']
                
                # 处理奇偶页不同
                if 'different_odd_even' in config:
                    section.odd_and_even_pages_header_footer = config['different_odd_even']
            
            print("  ✅ 更新页眉页脚")
            return True
            
        except Exception as e:
            print(f"  ❌ 页眉页脚设置失败: {e}")
            return False
    
    def _modify_header(self, header: Any, config: Dict[str, Any]) -> None:
        """修改页眉"""
        # 清除现有内容（如果需要）
        if config.get('clear_existing', False):
            for para in header.paragraphs:
                para.clear()
        
        # 添加新内容
        if 'text' in config:
            if not header.paragraphs:
                para = header.add_paragraph()
            else:
                para = header.paragraphs[0]
            
            para.text = config['text']
            
            # 应用样式
            if 'style' in config and config['style'] in self.doc.styles:
                para.style = config['style']
    
    def _modify_footer(self, footer: Any, config: Dict[str, Any]) -> None:
        """修改页脚"""
        # 类似于修改页眉
        if config.get('clear_existing', False):
            for para in footer.paragraphs:
                para.clear()
        
        if 'text' in config:
            if not footer.paragraphs:
                para = footer.add_paragraph()
            else:
                para = footer.paragraphs[0]
            
            para.text = config['text']
            
            if 'style' in config and config['style'] in self.doc.styles:
                para.style = config['style']
    
    def _apply_selective_modifications(self) -> None:
        """应用选择性修改"""
        # 根据配置选择性地修改文档部分
        for mod in self.selective_mods.get('modifications', []):
            target = mod.get('target')
            action = mod.get('action')
            
            if target == 'paragraphs':
                self._selective_modify_paragraphs(mod)
            elif target == 'tables':
                self._selective_modify_tables(mod)
            elif target == 'images':
                self._selective_modify_images(mod)
    
    def _selective_modify_paragraphs(self, mod: Dict[str, Any]) -> None:
        """选择性修改段落"""
        criteria = mod.get('criteria', {})
        actions = mod.get('actions', {})
        
        for para in self.doc.paragraphs:
            if self._match_paragraph_criteria(para, criteria):
                self._apply_paragraph_actions(para, actions)
    
    def _match_paragraph_criteria(self, para: Any, criteria: Dict[str, Any]) -> bool:
        """匹配段落条件"""
        # 样式匹配
        if 'style' in criteria:
            if para.style.name != criteria['style']:
                return False
        
        # 文本匹配
        if 'text_contains' in criteria:
            if criteria['text_contains'] not in para.text:
                return False
        
        # 正则匹配
        if 'text_regex' in criteria:
            if not re.search(criteria['text_regex'], para.text):
                return False
        
        return True
    
    def _apply_paragraph_actions(self, para: Any, actions: Dict[str, Any]) -> None:
        """应用段落操作"""
        # 修改样式
        if 'set_style' in actions:
            if actions['set_style'] in self.doc.styles:
                para.style = actions['set_style']
        
        # 修改文本
        if 'replace_text' in actions:
            old_text = actions['replace_text'].get('old')
            new_text = actions['replace_text'].get('new')
            if old_text and new_text:
                para.text = para.text.replace(old_text, new_text)
        
        # 添加前缀/后缀
        if 'add_prefix' in actions:
            para.text = actions['add_prefix'] + para.text
        
        if 'add_suffix' in actions:
            para.text = para.text + actions['add_suffix']
    
    def _selective_modify_tables(self, mod: Dict[str, Any]) -> None:
        """选择性修改表格"""
        # 实现表格的选择性修改
        pass
    
    def _selective_modify_images(self, mod: Dict[str, Any]) -> None:
        """选择性修改图片"""
        # 实现图片的选择性修改
        pass

def main():
    """测试结构修改器"""
    
    # 创建测试文档
    doc = Document()
    doc.add_heading('文档标题', 0)
    doc.add_heading('第一章', 1)
    doc.add_paragraph('第一章内容')
    doc.add_heading('第二章', 1)
    doc.add_paragraph('第二章内容')
    doc.add_heading('第2.1节', 2)
    doc.add_paragraph('第2.1节内容')
    
    # 创建测试配置
    import json
    from config_loader import ModifyConfigLoader
    
    config = {
        "modify_mode": "merge",
        "section_breaks": {
            "sections": [
                {"after_paragraph": 2, "type": "new_page"},
                {"after_paragraph": 4, "type": "odd_page"}
            ]
        },
        "document_structure": {
            "heading_levels": {
                "rules": [
                    {"from_level": 2, "new_level": 3, "text_pattern": "第\\d+\\.\\d+节"}
                ]
            }
        },
        "page_numbering": {
            "format": "decimal",
            "start_number": 1,
            "position": "footer_center"
        }
    }
    
    # 保存配置
    config_path = "test_structure_config.json"
    with open(config_path, 'w', encoding='utf-8') as f:
        json.dump(config, f, ensure_ascii=False, indent=2)
    
    # 测试修改器
    loader = ModifyConfigLoader(config_path)
    modifier = StructureModifier(doc, loader)
    
    results = modifier.modify_document_structure()
    
    print(f"\n修改结果:")
    print(f"分节符修改: {results['sections_modified']}个")
    print(f"标题更新: {results['headings_updated']}个")
    print(f"错误: {results['errors']}")
    
    # 保存文档
    doc.save("test_structured_document.docx")
    print(f"\n✅ 测试文档已保存")
    
    # 清理
    import os
    os.remove(config_path)
    os.remove("test_structured_document.docx")

if __name__ == "__main__":
    main()